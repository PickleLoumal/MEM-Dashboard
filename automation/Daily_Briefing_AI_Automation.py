"""
Daily Briefing AI Automation

Generates professional daily investment briefing reports using Perplexity AI.
Reads data from Google Sheets and uploads output to Google Drive.

Environment Variables Required:
- PERPLEXITY_API_KEY: Perplexity AI API key
- PERPLEXITY_API_URL: Perplexity API endpoint (optional, has default)
- GOOGLE_SHEETS_SPREADSHEET_ID: Google Sheets document ID
- GOOGLE_SHEETS_CREDENTIALS_FILE: Service account JSON file name
- GOOGLE_DRIVE_FOLDER_ID: Google Drive folder ID for output
- GOOGLE_OAUTH_CREDENTIALS_FILE: OAuth credentials file name
- GOOGLE_OAUTH_TOKEN_FILE: OAuth token cache file name
"""

import os
import datetime
import time
import re
import requests
import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import gspread
from google.oauth2.service_account import Credentials
from google.oauth2.credentials import Credentials as UserCredentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
import pickle

# Load environment variables from .env file if python-dotenv is available
try:
    from dotenv import load_dotenv
    # Load from project root .env file
    env_path = Path(__file__).resolve().parents[1] / ".env"
    load_dotenv(env_path)
except ImportError:
    pass  # python-dotenv not installed, rely on system environment variables

# =============================================================================
# Configuration from Environment Variables
# =============================================================================

# Perplexity API Configuration
PERPLEXITY_API_KEY = os.getenv("PERPLEXITY_API_KEY")
PERPLEXITY_API_URL = os.getenv("PERPLEXITY_API_URL", "https://api.perplexity.ai/chat/completions")

# Google Sheets Configuration
SPREADSHEET_ID = os.getenv("GOOGLE_SHEETS_SPREADSHEET_ID")
CREDENTIALS_FILE = os.getenv("GOOGLE_SHEETS_CREDENTIALS_FILE", "boxwood-sandbox-483506-u5-d246bd69f7d7.json")

# Google Drive Configuration
GOOGLE_DRIVE_FOLDER_ID = os.getenv("GOOGLE_DRIVE_FOLDER_ID")
OAUTH_CREDENTIALS_FILE = os.getenv("GOOGLE_OAUTH_CREDENTIALS_FILE", "oauth_credentials.json")
OAUTH_TOKEN_FILE = os.getenv("GOOGLE_OAUTH_TOKEN_FILE", "oauth_token.pickle")

# Validate required environment variables
def _validate_config():
    """Validate that required environment variables are set."""
    missing = []
    if not PERPLEXITY_API_KEY:
        missing.append("PERPLEXITY_API_KEY")
    if not SPREADSHEET_ID:
        missing.append("GOOGLE_SHEETS_SPREADSHEET_ID")
    if missing:
        raise EnvironmentError(
            f"Missing required environment variables: {', '.join(missing)}\n"
            f"Please set them in your .env file or system environment."
        )

def convert_markdown_to_word(markdown_text, doc):
    """
    将Markdown文本转换为Word文档格式，严格按照markdown语法标准
    支持红绿灯emoji表情: 🟢, 🟡, 🔴
    """
    lines = markdown_text.split('\n')
    in_code_block = False
    code_language = ""
    in_table = False
    table_obj = None

    for line in lines:
        # 处理代码块
        if line.strip().startswith('```'):
            if not in_code_block:
                # 开始代码块
                in_code_block = True
                code_language = line.strip()[3:].strip()
                continue
            else:
                # 结束代码块
                in_code_block = False
                code_language = ""
                continue

        if in_code_block:
            # 代码块内容用等宽字体，保持原始格式
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            continue

        # 检测表格（简单的markdown表格检测）
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                # 开始新表格
                in_table = True
                # 解析表头
                cells = [cell.strip() for cell in line.strip().split('|')[1:-1]]
                table_obj = doc.add_table(rows=1, cols=len(cells))
                table_obj.style = 'Table Grid'
                for i, cell_text in enumerate(cells):
                    table_obj.rows[0].cells[i].text = cell_text
            elif line.strip().replace('|', '').replace('-', '').replace(' ', '').replace(':', '') == '':
                # 表格分隔线，跳过
                continue
            else:
                # 添加表格行
                cells = [cell.strip() for cell in line.strip().split('|')[1:-1]]
                row_cells = table_obj.add_row().cells
                for i, cell_text in enumerate(cells):
                    if i < len(row_cells):
                        row_cells[i].text = cell_text
            continue
        else:
            if in_table:
                in_table = False
                table_obj = None

        # 跳过空行但保留段落间距
        if not line.strip():
            doc.add_paragraph()
            continue

        # 处理各级标题
        if line.startswith('# ') and not line.startswith('## '):
            heading = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## ') and not line.startswith('### '):
            heading = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### ') and not line.startswith('#### '):
            heading = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### ') and not line.startswith('##### '):
            heading = doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('##### ') and not line.startswith('###### '):
            heading = doc.add_heading(line[6:].strip(), level=5)

        # 处理无序列表
        elif re.match(r'^[ \t]*[-*+] ', line):
            indent_match = re.match(r'^([ \t]*)', line)
            indent_level = len(indent_match.group(1).expandtabs(4)) // 4 if indent_match else 0
            bullet_text = re.sub(r'^[ \t]*[-*+] ', '', line)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, bullet_text)

        # 处理有序列表
        elif re.match(r'^[ \t]*\d+\. ', line):
            bullet_text = re.sub(r'^[ \t]*\d+\. ', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, bullet_text)

        # 处理水平分割线
        elif line.strip() in ['---', '***', '___'] or re.match(r'^[ \t]*[-*_]{3,}[ \t]*$', line):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run('─' * 50)
            run.font.color.rgb = RGBColor(128, 128, 128)

        # 处理引用块
        elif line.strip().startswith('> '):
            quote_text = line.strip()[2:]
            p = doc.add_paragraph()
            p.left_indent = Inches(0.5)
            p.paragraph_format.left_indent = Inches(0.5)
            add_formatted_text(p, quote_text)

        # 处理普通段落
        else:
            if line.strip():
                p = doc.add_paragraph()
                add_formatted_text(p, line.strip())

    return doc

def add_formatted_text(paragraph, text):
    """向段落添加格式化文本，支持加粗、斜体、行内代码和emoji"""
    code_parts = re.split(r'(`[^`]*`)', text)

    for code_part in code_parts:
        if code_part.startswith('`') and code_part.endswith('`') and len(code_part) >= 2:
            run = paragraph.add_run(code_part[1:-1])
            run.font.name = 'Courier New'
        else:
            bold_parts = re.split(r'(\*\*[^*]*\*\*)', code_part)

            for bold_part in bold_parts:
                if bold_part.startswith('**') and bold_part.endswith('**') and len(bold_part) >= 4:
                    bold_content = bold_part[2:-2]
                    italic_parts = re.split(r'(\*[^*]+\*)', bold_content)

                    for italic_part in italic_parts:
                        if italic_part.startswith('*') and italic_part.endswith('*') and len(italic_part) >= 3:
                            run = paragraph.add_run(italic_part[1:-1])
                            run.bold = True
                            run.italic = True
                        else:
                            if italic_part:
                                run = paragraph.add_run(italic_part)
                                run.bold = True
                else:
                    italic_parts = re.split(r'(\*[^*]+\*)', bold_part)

                    for italic_part in italic_parts:
                        if italic_part.startswith('*') and italic_part.endswith('*') and len(italic_part) >= 3:
                            run = paragraph.add_run(italic_part[1:-1])
                            run.italic = True
                        else:
                            if italic_part:
                                paragraph.add_run(italic_part)

def authenticate_google_sheets():
    """认证 Google Sheets 并获取数据"""
    try:
        print("🔐 正在认证 Google Sheets...")

        scopes = [
            'https://www.googleapis.com/auth/spreadsheets',
            'https://www.googleapis.com/auth/drive'
        ]

        creds = Credentials.from_service_account_file(
            CREDENTIALS_FILE,
            scopes=scopes
        )

        client = gspread.authorize(creds)
        sheet = client.open_by_key(SPREADSHEET_ID).sheet1

        print("✅ Google Sheets 认证成功！")

        # 获取所有数据作为参考
        print("📊 正在读取表格数据...")
        all_values = sheet.get_all_values()

        # 格式化数据为文本
        sheet_data = "Google Sheets Data Summary:\n\n"
        for i, row in enumerate(all_values[:50], 1):  # 限制前50行避免过长
            sheet_data += f"Row {i}: {' | '.join([str(cell) for cell in row])}\n"

        print(f"✅ 成功读取 {len(all_values)} 行数据")

        return sheet_data

    except Exception as e:
        print(f"❌ Google Sheets 认证或读取失败: {str(e)}")
        return None

def get_google_drive_credentials():
    """
    获取 Google Drive OAuth 2.0 凭据
    首次运行会打开浏览器授权，之后使用缓存的 token
    """
    creds = None
    script_dir = os.path.dirname(os.path.abspath(__file__))
    token_path = os.path.join(script_dir, OAUTH_TOKEN_FILE)
    oauth_path = os.path.join(script_dir, OAUTH_CREDENTIALS_FILE)

    # 尝试加载已保存的 token
    if os.path.exists(token_path):
        with open(token_path, 'rb') as token:
            creds = pickle.load(token)

    # 如果没有有效凭据，进行授权
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            print("🔄 刷新 Google Drive 授权...")
            creds.refresh(Request())
        else:
            print("🔐 首次授权：将打开浏览器进行 Google 账号授权...")
            flow = InstalledAppFlow.from_client_secrets_file(
                oauth_path,
                scopes=['https://www.googleapis.com/auth/drive.file']
            )
            creds = flow.run_local_server(port=0)
            print("✅ 授权成功！")

        # 保存 token 供下次使用
        with open(token_path, 'wb') as token:
            pickle.dump(creds, token)

    return creds

def upload_to_google_drive(file_path, folder_id=None):
    """
    上传文件到 Google Drive（使用 OAuth 2.0 用户认证）

    Args:
        file_path: 本地文件路径
        folder_id: Google Drive 目标文件夹 ID（可选）

    Returns:
        dict: 包含文件 ID 和链接，失败返回 None
    """
    try:
        print(f"\n📤 正在上传文件到 Google Drive: {os.path.basename(file_path)}")

        # 获取 OAuth 凭据
        creds = get_google_drive_credentials()

        # 创建 Drive 服务
        service = build('drive', 'v3', credentials=creds)

        # 文件元数据
        file_metadata = {
            'name': os.path.basename(file_path),
        }

        # 如果指定了文件夹 ID，设置父目录
        if folder_id:
            file_metadata['parents'] = [folder_id]

        # 上传文件
        media = MediaFileUpload(
            file_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            resumable=True
        )

        file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, name, webViewLink'
        ).execute()

        print(f"   ✅ 上传成功！")
        print(f"   📄 文件名: {file.get('name')}")
        print(f"   🔗 文件 ID: {file.get('id')}")
        if file.get('webViewLink'):
            print(f"   🌐 查看链接: {file.get('webViewLink')}")

        return file

    except Exception as e:
        print(f"   ❌ 上传失败: {str(e)}")
        return None

def call_perplexity_api(prompt, max_retries=3):
    """调用 Perplexity Sonar Deep Research API"""
    headers = {
        "Authorization": f"Bearer {PERPLEXITY_API_KEY}",
        "Content-Type": "application/json"
    }

    data = {
        "model": "sonar-deep-research",
        "messages": [
            {
                "role": "system",
                "content": "You are a highly intelligent AI financial analyst with deep research capabilities, tasked with generating professional Daily Briefing reports for institutional investment decision-making."
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    }

    for attempt in range(max_retries):
        try:
            print(f"🔄 尝试第 {attempt + 1}/{max_retries} 次调用 Perplexity API...")

            response = requests.post(
                PERPLEXITY_API_URL,
                headers=headers,
                json=data,
                timeout=3600  # 1小时超时
            )
            response.raise_for_status()
            result = response.json()

            if result and "choices" in result and len(result["choices"]) > 0:
                content = result["choices"][0]["message"]["content"]
                # 移除 <think> 标签内容
                content = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL)
                return content.strip()
            else:
                raise Exception("API返回格式异常")

        except Exception as e:
            print(f"⚠️  API调用失败 (尝试 {attempt + 1}/{max_retries}): {str(e)}")
            if attempt < max_retries - 1:
                time.sleep(10)
            else:
                raise

    return None

# Daily Briefing Prompt Template
daily_briefing_prompt = """Today is Jan 8, 2026, Hong Kong time, and Jan 7, 2025,ET Time
Attachments(Includes latest prices & index, Briefing.com Page One, Stock & Bond latest trends, Barrons latest news, Valueline last day conclusion): https://docs.google.com/spreadsheets/d/1cS4iSRck__DHXh1U7PfqfSmVxc_Aa_Tw2qbs35wI4i4/edit?gid=0#gid=0 In this we have latest price and latest news. You are an AI financial analyst tasked with generating a professional Daily Briefing report for institutional investment decision-making. The report must be data-driven, forward-looking, concise, and highly actionable. Maintain a formal, professional tone with clear, well-structured sentences and precise language. Use bullet points and tables where they improve clarity and efficiency.
All data must be accurate and current as of the report date and must come from sources within the last 1 month. If any data older than 1 month is used, explicitly add a footnote explaining why it is still relevant and cannot yet be updated. Avoid vague time references such as "YTD (July)"; always use precise date ranges (for example, "YTD through December 11, 2025"). Do not hallucinate data and do not reuse stale data from prior reports.
Primarily analyze from the attached document. You may also fetch incremental or verifying data from reputable sources such as Briefing.com, Reuters, Bloomberg, FactSet, Federal Reserve, Trading Economics, FRED, Nasdaq, and others. For each numerical or factual data point, cite its source and exact timestamp inline in the text (for example, "[Briefing.com, 11-Dec-25 16:30 ET"). Be skeptical and data-driven, explicitly flagging any inconsistencies across sources or within the attachments.
Use tools such as: web_search (for verification, for example, "S&P 500 closing value Dec 11, 2025 site:reuters.com"), web_search_with_snippets (for quick fact checks), and code_execution (for calculations such as totals, averages, percentage changes, or simple trend analyses). Do not hallucinate technical indicators; compute or infer them only when data is available. Ensure strict political neutrality and substantiate any non-trivial claim with data.
The Daily Briefing shall serve the following purposes:
Update on the latest market news and changes since the last meeting (Dec 11, 2025).
Monitor risks and implications for existing portfolio exposures and positions.
Generate new investment and trading ideas with clear, actionable implementation details.
Structure the report exactly as follows:
Title block (fixed pattern):
"Daily Briefing Report
Report Date: [Day of week], [Month] [Day], [Year], [Time] HKT
Market Status: [U.S. markets status]; [Asia-Pacific session status]; [Europe status]
Portfolio Stance: [Market regime description] — [Key probability or data point, e.g., 'Fed rate cut probability >87% for December 9–10 FOMC meeting; risk-on sentiment supported by easing expectations but tempered by weak manufacturing data and tariff uncertainty']"
Executive Summary (Concise Version)
Begin the report with a concise, approximately half-page Executive Summary. Use bullets for clarity.
Include:
3–5 bullet points on key market moves since the last meeting (major indices, yields, FX, commodities, crypto, notable macro events).
3–5 bullet points on top risks (policy risk, geopolitical risk, sector-specific risks such as healthcare cost inflation, China growth concerns, etc.).
2–4 bullet points on immediate actions for the portfolio (for example, "Reduce or hedge UNH exposure due to accelerating medical cost trends and sector underperformance").
Add 1–2 sentences on implications of key macro/political events for specific sectors, explicitly including examples such as how Trump tariff rhetoric or trade policy may impact industrials, exporters, or specific sectors (for example, industrials, consumer discretionary).
The Executive Summary must be tightly aligned with later sections and must not introduce contradictions.
News and Market Briefing: Update on Changes Since the Last Meeting
Provide a structured update on changes since the prior meeting date (Dec 9, 2025). Cover the following categories:
Public Equities (focus on US, China/Hong Kong, Europe, Japan, UK, and other notable markets; include sector rotation and market breadth).
Fixed Income (yield curve shape, especially UST 2-yr and 10-yr levels and changes, curve steepening/flattening).
Currencies & Policy (EURUSD, USDJPY, DXY; key central bank updates; macroeconomic data; politics/geopolitics).
Commodities (e.g., WTI crude, gold, silver, copper; clearly specify date ranges such as "YTD through [date]").
Cryptocurrency Market (e.g., BTC, ETH movements and volatility regimes).
Private Investments (trends in PE/VC, hedge funds, private credit and infrastructure).
For each category:
Provide 2–3 headlines.
Summarize trends/patterns.
Add relevant context and implications (why the development matters for markets, risk appetite, and/or the portfolio).
1.1 Market Snapshot Table
Start the section with a Market Snapshot table.
Columns:
Index/Item
Closing Value
Change (%) or bp
1-Week Technical Indicators
1-Week Trend Identification
Why It Matters
Source
Technical requirements:
For each line item, add a brief 1-week technical analysis comment:
Include basic indicators such as 1-week price change, relative performance vs peers/benchmark, and simple signals such as short-term moving average crossovers, overbought/oversold conditions, or consolidation ranges if supported by the data.
Identify whether the short-term trend is "Uptrend", "Downtrend", or "Consolidation", and briefly highlight any significant technical insights (for example, "testing prior resistance", "breakout above recent range", "bearish divergence vs breadth").
Keep the commentary concise but specific.
1.2 Sector Rotation Analysis (Deepened and Expanded)
Produce a multi-temporal, rigorous sector rotation analysis for major equity sectors (e.g., Technology, Communication Services, Consumer Discretionary, Consumer Staples, Health Care, Financials, Industrials, Energy, Materials, Real Estate, Utilities). This section must explicitly incorporate:

Part 1 – Trailing 12-Month Sector Performance Trends and Key Drivers
Provide a table with columns such as: Sector | Approx. Trailing 12M Performance (%) | Direction (Up/Down/Mixed) | Key Drivers (macro, policy, demand cycles) | 1-Week Technical View.
Use available data from Briefing.com's "Stock Market Update" and other authoritative sources. Where exact 12-month performance data is not available from the attachments, classify trends qualitatively (for example, "structural leader", "modestly positive", "mixed", "laggard"), and label them as directional views rather than precise numbers.
Highlight major drivers (for example, AI capex for Technology, energy prices and capex cycles for Energy, medical cost inflation for Health Care, etc.).
Incorporate short-term (1-week) technical signals into this table or a parallel column where possible.

Part 2 – Consensus Outlook for Sector Rotations and Main Reasons
Summarize consensus expectations for medium-term sector rotations (for example, structural shift toward growth/AI/energy infrastructure vs. defensives).
Reference factors such as valuation, earnings revision trends, positioning, policy, and macro regime.
Explicitly note where consensus appears crowded or subject to key risks.

Part 3 – Recent Performance: Most Recent Session and Last Five Sessions
Summarize sector performance for the most recent trading session and the last 5 sessions.
First list leading sectors with positive performance (largest positive percentage changes at the top).
Then list lagging sectors with negative performance.
Comment concisely on whether short-term patterns suggest rotation, consolidation, or conflicting signals (for example, "energy leading despite lower crude price", "healthcare underperforming due to cost inflation headlines").

Part 4 – Broader Implications
Explain the implications of these sector rotations for:
Other industries and factor exposures.
Geopolitics and macro regimes (for example, defense spending, trade tensions, China growth).
Global and national economies (for example, late-cycle rotation into defensives, commodities, or infrastructure).
Explicitly link this to portfolio sectors, such as healthcare (UNH), energy/miners (GDX, GDXJ, URA), technology/AI (TSM, other Magnificent 7 exposures via infrastructure), etc.
"Why it matters" statement:
Immediately after the integrated sector rotation discussion, add a short paragraph starting with "Why it matters:" providing human-like insight and skepticism. For example:
"Why it matters: the market appears to be leaning back toward growth and consumer cyclicals even as the policy path remains uncertain—this may prove fragile given recent inflation data and the structural support for energy and industrials over the trailing 12 months."
Explicitly question inconsistencies or fragile narratives when warranted (for example, "bond yields eased despite an uptick in inflation—possibly reflecting policy uncertainty or growth concerns rather than a clean disinflation signal").
Data Verification Table (end of Sector Rotation section):
Add a compact table titled "Sector Rotation Data Verification".
Columns:
Data Point
Source 1
Source 2
Source 3
Notes
Use this table to show cross-verification for key sector performance and index data points, ideally referencing at least three authoritative sources (for example, Briefing.com, Reuters, Bloomberg, FactSet). Note any data that is only verified by a single source, and flag where data is directional/approximate due to incomplete disclosure in the attachments.
"20 small points per sector" requirement:
To satisfy the "20 key points/little things each sector" expectation, ensure that across Parts 1–4 you provide granular observations per sector (for example, subsector divergences, leadership within the sector, earnings revisions, flows/ETF positioning, valuation comments, policy/ regulatory headlines, and cross-market links). These do not need to be numbered one by one, but the commentary should be sufficiently granular and multi-angle so that each sector has deep coverage rather than a single high-level sentence.
1.3 Fixed Income (Yield Curve and Rates)
Describe yield curve shape (e.g., 2s/10s spread, steepening/flattening).
Provide current 2-year and 10-year UST yields and relevant changes (in bps).
Add 1-week technical assessment: has the curve steepened or flattened, have yields broken recent ranges, etc.
Summarize key drivers (data surprises, central bank communication, inflation expectations, auctions).
Discuss implications for rate-sensitive sectors (real estate, utilities, growth equities) and for the portfolio.
1.4 Currencies & Policy
Cover DXY, EURUSD, USDJPY and key central bank updates (Fed, ECB, BoE, BoJ, RBA, etc.).
Consolidate central bank updates into a single table to avoid redundancy.
Highlight major policy meetings, SEP updates, and any shifts in forward guidance or balance of risks.
Discuss implications for global risk appetite and specifically for portfolio exposures such as BABA, 2840.HK, TSM (FX translation and policy risk).
1.5 Commodities (with Corrected Data and Clear Date Ranges)
Provide concise updates on key commodities (e.g., WTI, Brent if relevant, gold, silver, copper, uranium if relevant).
Correct any previous vague references such as "YTD through July" to precise ranges like "YTD through December 11, 2025".
Ensure data for indices like the NYSE Arca Gold Miners Index (GDMNTR) and silver performance are updated to the current date and correctly labeled (for example, "+X% YTD through December 11, 2025").
Clarify data source and update frequency (for example, daily close, end-of-month, etc.).
Add 1-week technical observations (for example, support/resistance levels, short-term momentum).
1.6 Cryptocurrency Market
Summarize BTC and ETH price levels, percentage changes, and volatility regimes.
Highlight key ranges (for example, consolidation below a round-number resistance), and any correlation shifts vs. risk assets or real yields.
Note that Bitcoin content may need integration of comments from a specific team member (e.g., "Man On") when those are provided; until then, focus on neutral, data-driven observations.
Link implications to portfolio crypto exposures (e.g., ETHA, IBIT).
1.7 Private Investments
Summarize recent trends in PE/VC, private credit, infrastructure, and continuation funds, emphasizing data within the last month.
Highlight shifts in deal volume, sector focus (especially infrastructure, technology, financial services), and exit environment.
Discuss how these trends affect the portfolio's private funds.
1.8 Magnificent 7 Infrastructure Investment Section
Add a dedicated subsection analyzing Magnificent 7–related infrastructure investments:
Data center build-out (capacity, capex, geographic diversification).
AI computing capacity and semiconductor supply chain implications.
Related listed opportunities (data center REITs, equipment vendors, power and grid infrastructure, etc.).
Identify key news, opportunities, and risks, with clear connections to holdings such as TSM and to potential idea generation in Section 3.
1.9 Heatmaps from Finviz
Add two visual references (described if embedding is not possible):
A 1-day Finviz heatmap for major US sectors and names.
A 3-month Finviz heatmap to show medium-term relative performance.
Place these descriptions or images within the Market Briefing or Sector Rotation section and briefly interpret their key messages (for example, concentration of gains in mega-cap tech vs. broader participation).
1.10 SPY Live Market Chart Description
End the News and Market Briefing section with a concise description of a live SPY (S&P 500 ETF) chart or a placeholder description, covering:
Short-term price action vs. recent highs/lows.
Volume characteristics.
Key technical levels and indicators (for example, moving averages, Bollinger Bands).
Likely inflection points around major events (for example, FOMC).
Portfolio Adjustments & Actions
Analyze how the latest news and market movements affect existing portfolio exposures and positions. The portfolio includes (but is not limited to):
Funds:
JP Morgan Infrastructure Investments Fund (IIF)
JP Morgan Q Systematic Private Investors Offshore Ltd.
JP Morgan Uncorrelated Hedge Fund Strategies S.A. SICAV-RAIF
BNP Paribas HarbourVest's Dover Street XI AIF SCSp (Dover Street XI or Fund XI)
BNP Paribas CIO Strategy Fund
Global Access Strategies SPC Ltd.
JP Morgan Vintage 2024 Private Investments Offshore SICAV-RAIF S.C.Sp.
Ardian Access, SCA., SICAV-RAIF
Janus Henderson Horizon Biotechnology Fund
iCapital Millennium Fund II, Ltd.
DNCA Invest-Alpha Bonds-A-ACC-EUR
Equities/ETFs:
BA, BABA, BRK-B, ETHA, GDX, GDXJ, GLD, GWMGF, IBIT, INTC, MAAFF, SLV, TSM, UNH, URA, 2840.HK (and any others present in the attachment).
2.1 Portfolio Monitoring List (Not Table)
Provide a list (not a table) where each line or block contains at least:
Ticker/Fund
Close
Today Change (%)
Traffic-light flag: 🟢 Stable/Hold, 🟡 Watch/Maintain, 🔴 Reduce/Hedge
Driver (main fundamental/technical/macro driver)
Latest News (if available) and sentiment
Source
2.2 Improved Portfolio Recommendation Logic (Forward-Looking)
In line with the leadership's instructions:
Red alert "Reduce" or "Hedge" recommendations must be based on the future outlook, not merely on recent poor price performance.
Every recommendation must explicitly discuss forward-looking drivers, including earnings trajectory, sector fundamentals, macro/policy regime, valuations, and risk/reward.
Provide explicit risk/reward scenarios where relevant (for example, base/bull/bear or upside vs. downside range).
Cross-reference portfolio actions with the market sections. For example:
"UNH: 🔴 Reduce/Hedge — consistent with Health Care sector underperformance and worsening medical cost inflation trends discussed in the Sector Rotation section."
"BABA: 🟡 Watch/Maintain — aligns with continued China/HK weakness, FX and policy risks as described under China/HK markets and Currencies & Policy."
2.3 Portfolio Stance & Regime Assessment
Summarize the portfolio's overall stance in the current market regime, for example, "Late-cycle disinflation with hawkish pause and elevated policy uncertainty".
Discuss:
Growth momentum (for example, "moderating but resilient").
Inflation trends vs. target.
Policy stance (for example, "hawkish cut", "pause with asymmetric risks").
Valuations (index-level, major sectors).
Volatility and credit conditions.
Key portfolio-level risks (Fed policy error, geopolitical escalation, China slowdown, healthcare cost inflation, FX volatility).
Portfolio positioning (e.g., quality/AI tilt, precious metals overweight, underweight financials/healthcare).
Explicit recommended stance (for example, "cautiously constructive with a defensive tilt") and immediate actions (for example, reduce UNH, tighten stops on BABA/BA/INTC, maintain or add on dips to GDX/GDXJ/TSM, etc.).
Market Opportunities & Tradable Ideas
Distill actionable investment and trading ideas from the news update and portfolio analysis.
Provide 1–2 ideas for each horizon:
Short-term (0–3 months).
Mid-term (3–12 months).
Long-term (>12 months).
For each idea, present the content as rows/blocks (not in a table), with the following fields clearly labeled:
Idea ID
Asset / Direction / Horizon
Thesis (concise but rich reasoning, referencing data, trends, sector dynamics, and cross-source consistency or inconsistencies; incorporate human-like insight and skepticism).
Catalysts (with specific dates or date ranges).
Implementation (for example, position sizing, instrument choice, options structure if relevant).
Entry / Target / Stop (include an explicit risk–reward ratio).
Projected Return / Risk (qualitative plus approximate percentage ranges if supported by data).
Risks (what can invalidate the thesis).
Monitoring (what indicators, data releases, price levels, or flows to watch).
Use bullets inside each block for clarity, and bold key actions (for example, "Enter on pullback to [level]", "Tighten stop to [level] after event").
3.1 Idea Tracking Table
At the end of the ideas section, include an Idea Tracking table with columns:
ID
Date
Asset
Direction
Entry
Target
Stop
Next Review
If no prior ideas are provided, assume the current set is the first cohort and initialize the tracking table accordingly. When prior ideas exist, briefly remark on their viability based on the latest data (for example, "still in play", "stopped out", "reached first target").
Appendix: Sources & Timestamps
List all sources and timestamps cited in the report, including:
Web and data sources (for example, "Briefing.com 'Stock Market Update' (11-Dec-25 16:30 ET)").
Attachment-derived data, clearly labeled as manually collected (for example, "Portfolio price data from manually compiled file based on [data provider] as of [timestamp]").
Formatting and Consistency Requirements
Use bold formatting for key actions, key numbers, and important risk/implication phrases to enhance readability.
Use bullet points and tables where they help structure complex information (comparisons, enumerations, data presentation).
Ensure the context is coherent across all sections, with no analytical contradictions. When contradictions in the provided documents or between sources are identified, explicitly highlight them and explain how they are treated in the analysis (for example, "Source A reports X, Source B reports Y; base-case view leans toward X due to [reason]").
Maintain strict internal consistency between the Executive Summary, the detailed sections, and the recommended portfolio actions.
At the end of the generation process, internally validate:
All data is within the last 1 month or explicitly footnoted if older.
Every data table that contains market/sector/asset data includes a 1-week technical view and trend identification.
Two Finviz heatmaps (1-day and 3-month) are incorporated and interpreted.
The Executive Summary contains clear implications (including tariff/trade policy sector impact).
Key drivers are analyzed with dynamic trend context (historical trajectory, strengthening/weakening, potential inflection points, and forward direction).
Commodity data uses precise date ranges and current figures.
Sector rotation analysis is deep, multi-temporal, and supported by a Data Verification Table.
Portfolio "Reduce" recommendations are grounded in forward-looking analysis and explicitly linked back to sector/macro conclusions."""

########################## 主程序 ##########################

def main():
    # Validate configuration before starting
    _validate_config()

    start_time = datetime.datetime.now()
    today = start_time.strftime("%Y-%m-%d")
    today_display = start_time.strftime("%B %d, %Y")

    print("=" * 70)
    print("Daily Briefing AI Automation")
    print("=" * 70)
    print(f"📅 日期: {today_display}")
    print(f"⏰ 开始时间: {start_time.strftime('%H:%M:%S')}")
    print()

    # 创建输出目录（基于脚本所在目录）
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_base_dir = os.path.join(script_dir, "output")
    output_dir = os.path.join(output_base_dir, today)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
        print(f"✅ 创建输出目录: {output_dir}\n")

    # 1. 读取 Google Sheets 数据
    sheet_data = authenticate_google_sheets()

    if not sheet_data:
        print("⚠️  无法读取 Google Sheets 数据，将继续使用 prompt 中的 URL 引用")
        sheet_context = ""
    else:
        sheet_context = f"\n\n[Additional Context from Google Sheets]:\n{sheet_data}\n"

    # 2. 组合完整 prompt
    full_prompt = daily_briefing_prompt + sheet_context

    # 3. 调用 Perplexity API
    try:
        print("\n🔄 开始调用 Perplexity Sonar Deep Research API...")
        print("⏳ 这可能需要几分钟时间，因为 AI 正在进行深度研究和数据搜集...\n")

        response_content = call_perplexity_api(full_prompt)

        if not response_content or len(response_content) < 500:
            print("❌ AI 返回内容为空或过短")
            return False

        print(f"✅ AI 分析完成！内容长度: {len(response_content)} 字符\n")

        # 4. 转换为 Word 文档 - 长版本
        print("📝 转换 Markdown 格式到 Word 文档 (长版本)...")
        doc_long = Document()
        doc_long = convert_markdown_to_word(response_content, doc_long)

        # 5. 保存长版本文档
        output_filename_long = os.path.join(output_dir, f"Daily Briefing {today}_Long Version.docx")
        doc_long.save(output_filename_long)

        long_end_time = datetime.datetime.now()
        long_processing_time = (long_end_time - start_time).total_seconds()

        print(f"✅ 长版本保存成功: {output_filename_long}")
        print(f"📄 内容长度: {len(response_content)} 字符")
        print(f"⏱️  处理耗时: {long_processing_time:.1f} 秒 ({long_processing_time/60:.1f} 分钟)")

        # 6. 生成快速版本（5页摘要）
        print("\n" + "=" * 70)
        print("📋 开始生成快速版本（5页摘要）...")
        print("=" * 70)

        summarize_prompt = f"""The current report in this attachment is {len(response_content)} characters long and is the detailed version. Could you please summarize it into 5 pages and convert it into the quick version?

Requirements for the Quick Version:
1. Maximum 5 pages when converted to Word document
2. Keep all critical information and key insights
3. Maintain the same structure but condense each section
4. Preserve all traffic-light indicators (🟢, 🟡, 🔴)
5. Keep the most important data points, charts, and tables
6. Focus on actionable insights and immediate portfolio implications
7. Use bullet points and concise language
8. Maintain professional tone

Original Report Content:
{response_content}

Please generate the condensed 5-page Quick Version now."""

        try:
            print("🔄 调用 Perplexity API 生成快速版本...")
            quick_response = call_perplexity_api(summarize_prompt)

            if not quick_response or len(quick_response) < 200:
                print("⚠️  快速版本生成失败，将只保留长版本")
            else:
                print(f"✅ 快速版本生成完成！内容长度: {len(quick_response)} 字符\n")

                # 7. 转换快速版本为 Word 文档
                print("📝 转换快速版本到 Word 文档...")
                doc_quick = Document()
                doc_quick = convert_markdown_to_word(quick_response, doc_quick)

                # 8. 保存快速版本文档
                output_filename_quick = os.path.join(output_dir, f"Daily Briefing {today}_Quick Version.docx")
                doc_quick.save(output_filename_quick)

                quick_end_time = datetime.datetime.now()
                quick_processing_time = (quick_end_time - long_end_time).total_seconds()
                total_time = (quick_end_time - start_time).total_seconds()

                print(f"✅ 快速版本保存成功: {output_filename_quick}")
                print(f"📄 内容长度: {len(quick_response)} 字符")
                print(f"⏱️  快速版本处理耗时: {quick_processing_time:.1f} 秒 ({quick_processing_time/60:.1f} 分钟)")
                print(f"⏱️  总耗时: {total_time:.1f} 秒 ({total_time/60:.1f} 分钟)")

                print("\n" + "=" * 70)
                print("🎉 Daily Briefing 生成完成！")
                print(f"📁 长版本: {output_filename_long}")
                print(f"📁 快速版本: {output_filename_quick}")
                print("=" * 70)

                # 9. 上传到 Google Drive（如果配置了文件夹 ID）
                if GOOGLE_DRIVE_FOLDER_ID:
                    print("\n" + "=" * 70)
                    print("☁️  开始上传文件到 Google Drive...")
                    print("=" * 70)

                    upload_to_google_drive(output_filename_long, GOOGLE_DRIVE_FOLDER_ID)
                    upload_to_google_drive(output_filename_quick, GOOGLE_DRIVE_FOLDER_ID)

                    print("\n" + "=" * 70)
                    print("✅ Google Drive 上传完成！")
                    print("=" * 70)

        except Exception as e:
            print(f"⚠️  快速版本生成出错: {str(e)}")
            print("✅ 长版本已成功保存")
            # 仍然尝试上传长版本
            if GOOGLE_DRIVE_FOLDER_ID:
                upload_to_google_drive(output_filename_long, GOOGLE_DRIVE_FOLDER_ID)

        return True

    except Exception as e:
        print(f"❌ 生成过程出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    main()
