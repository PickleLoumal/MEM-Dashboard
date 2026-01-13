"""
Forensic Accounting Automation

Generates forensic accounting analysis reports for listed companies using Perplexity AI.
Reads company list from Excel file and fetches stock data from Yahoo Finance.

Environment Variables Required:
- PERPLEXITY_API_KEY: Perplexity AI API key
- PERPLEXITY_API_URL: Perplexity API endpoint (optional, has default)
"""

import os
import datetime
import time
import re
import json
from pathlib import Path

import pandas as pd
import requests
import yfinance as yf
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Load environment variables from .env file if python-dotenv is available
try:
    from dotenv import load_dotenv
    env_path = Path(__file__).resolve().parents[1] / ".env"
    load_dotenv(env_path)
except ImportError:
    pass

# =============================================================================
# Configuration from Environment Variables
# =============================================================================

PERPLEXITY_API_KEY = os.getenv("PERPLEXITY_API_KEY")
PERPLEXITY_API_URL = os.getenv("PERPLEXITY_API_URL", "https://api.perplexity.ai/chat/completions")


def _validate_config():
    """Validate that required environment variables are set."""
    if not PERPLEXITY_API_KEY:
        raise EnvironmentError(
            "Missing required environment variable: PERPLEXITY_API_KEY\n"
            "Please set it in your .env file or system environment."
        )

class YahooFinanceDataFetcher:
    """
    Yahoo Finance 数据获取器，专门用于获取股价和市值数据
    完全免费，无需 API Key
    """

    def __init__(self):
        self.session = None

    def connect(self):
        """
        连接到 Yahoo Finance（实际上不需要连接，yfinance 是免费的）
        """
        try:
            print("🔌 连接到 Yahoo Finance API...")
            self.session = True
            print("✅ Yahoo Finance 已就绪（免费，无需 API Key）")
            return True
        except Exception as e:
            print(f"❌ Yahoo Finance 初始化失败: {e}")
            return False

    def get_stock_data(self, symbol):
        """
        获取股票的前一天收盘价和市值数据

        参数:
        symbol: 股票代码 (如 '0700.HK', 'AAPL', 'TSLA')

        返回:
        dict: {'last_price': 前一天收盘价, 'market_cap': 市值, 'currency': 货币}
        """
        try:
            if not self.session:
                print("请先连接 Yahoo Finance")
                return None

            print(f"📊 从 Yahoo Finance 获取 {symbol} 的数据...")

            # 使用 yfinance 获取股票数据
            ticker = yf.Ticker(symbol)

            # 获取股票信息
            info = ticker.info

            # 获取历史数据（最近2天，确保能拿到前收盘价）
            hist = ticker.history(period='2d')

            # 提取数据
            price_value = None
            market_cap_value = None
            currency_value = None

            # 获取前收盘价
            if not hist.empty and len(hist) > 0:
                # 获取最后一个交易日的收盘价（前收盘价）
                price_value = hist['Close'].iloc[-1]
                print(f"✅ 前收盘价: {price_value}")

            # 获取市值
            if 'marketCap' in info:
                market_cap_value = info['marketCap']
                print(f"✅ 市值: {market_cap_value}")

            # 获取货币
            if 'currency' in info:
                currency_value = info['currency']
                print(f"✅ 货币: {currency_value}")
            elif 'financialCurrency' in info:
                currency_value = info['financialCurrency']
                print(f"✅ 货币: {currency_value}")

            result = {
                'symbol': symbol,
                'last_price': price_value,
                'market_cap': market_cap_value,
                'currency': currency_value
            }

            print(f"✅ 数据获取成功: 前收盘价={result['last_price']}, 市值={result['market_cap']}, 货币={result['currency']}")
            return result
        except Exception as e:
            print(f"❌ 获取 {symbol} 数据时出错: {e}")
            print(f"💡 提示: 请确认股票代码格式正确（如 0700.HK, AAPL, TSLA）")
            return None

    def disconnect(self):
        """
        断开连接（Yahoo Finance 不需要断开）
        """
        try:
            if self.session:
                self.session = None
                print("✅ Yahoo Finance 会话已结束")
        except Exception as e:
            print(f"❌ 断开连接失败: {e}")

def convert_markdown_to_word(markdown_text, doc):
    """
    将Markdown文本转换为Word文档格式，严格按照markdown语法标准
    支持红绿灯emoji表情: 🟢, 🟡, 🔴
    """
    lines = markdown_text.split('\n')
    in_code_block = False
    code_language = ""
    in_table = False
    table_obj = None

    for line in lines:
        # 处理代码块
        if line.strip().startswith('```'):
            if not in_code_block:
                # 开始代码块
                in_code_block = True
                code_language = line.strip()[3:].strip()  # 获取语言标识
                continue
            else:
                # 结束代码块
                in_code_block = False
                code_language = ""
                continue

        if in_code_block:
            # 代码块内容用等宽字体，保持原始格式
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            continue

        # 检测表格（简单的markdown表格检测）
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                # 开始新表格
                in_table = True
                # 解析表头
                cells = [cell.strip() for cell in line.strip().split('|')[1:-1]]
                table_obj = doc.add_table(rows=1, cols=len(cells))
                table_obj.style = 'Table Grid'
                for i, cell_text in enumerate(cells):
                    table_obj.rows[0].cells[i].text = cell_text
            elif line.strip().replace('|', '').replace('-', '').replace(' ', '').replace(':', '') == '':
                # 表格分隔线，跳过
                continue
            else:
                # 添加表格行
                cells = [cell.strip() for cell in line.strip().split('|')[1:-1]]
                row_cells = table_obj.add_row().cells
                for i, cell_text in enumerate(cells):
                    if i < len(row_cells):
                        row_cells[i].text = cell_text
            continue
        else:
            if in_table:
                in_table = False
                table_obj = None

        # 跳过空行但保留段落间距
        if not line.strip():
            doc.add_paragraph()
            continue

        # 处理各级标题 - 必须是行首且后面有空格
        if line.startswith('# ') and not line.startswith('## '):
            heading = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## ') and not line.startswith('### '):
            heading = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### ') and not line.startswith('#### '):
            heading = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### ') and not line.startswith('##### '):
            heading = doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('##### ') and not line.startswith('###### '):
            heading = doc.add_heading(line[6:].strip(), level=5)

        # 处理无序列表 - 支持缩进
        elif re.match(r'^[ \t]*[-*+] ', line):
            # 计算缩进级别
            indent_match = re.match(r'^([ \t]*)', line)
            indent_level = len(indent_match.group(1).expandtabs(4)) // 4 if indent_match else 0

            # 提取列表内容
            bullet_text = re.sub(r'^[ \t]*[-*+] ', '', line)

            # 创建列表项
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, bullet_text)

        # 处理有序列表
        elif re.match(r'^[ \t]*\d+\. ', line):
            # 提取列表内容
            bullet_text = re.sub(r'^[ \t]*\d+\. ', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, bullet_text)

        # 处理水平分割线
        elif line.strip() in ['---', '***', '___'] or re.match(r'^[ \t]*[-*_]{3,}[ \t]*$', line):
            # 添加水平线
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run('─' * 50)
            run.font.color.rgb = RGBColor(128, 128, 128)

        # 处理引用块 (以 > 开头)
        elif line.strip().startswith('> '):
            quote_text = line.strip()[2:]
            p = doc.add_paragraph()
            p.left_indent = Inches(0.5)
            p.paragraph_format.left_indent = Inches(0.5)
            add_formatted_text(p, quote_text)

        # 处理普通段落
        else:
            if line.strip():
                p = doc.add_paragraph()
                add_formatted_text(p, line.strip())

    return doc

def add_formatted_text(paragraph, text):
    """
    向段落添加格式化文本，按照正确的markdown优先级处理
    支持emoji（包括红绿灯: 🟢, 🟡, 🔴）
    """
    # 按照markdown标准的优先级处理：先处理加粗，再处理斜体，最后处理行内代码

    # 1. 先处理行内代码 `code` (最高优先级，不能被其他格式影响)
    code_parts = re.split(r'(`[^`]*`)', text)

    for code_part in code_parts:
        if code_part.startswith('`') and code_part.endswith('`') and len(code_part) >= 2:
            # 行内代码 - 直接添加，不处理其他格式
            run = paragraph.add_run(code_part[1:-1])
            run.font.name = 'Courier New'
        else:
            # 2. 处理加粗文本 **text** (高优先级)
            bold_parts = re.split(r'(\*\*[^*]*\*\*)', code_part)

            for bold_part in bold_parts:
                if bold_part.startswith('**') and bold_part.endswith('**') and len(bold_part) >= 4:
                    # 加粗文本 - 内容还可能包含斜体
                    bold_content = bold_part[2:-2]
                    # 在加粗文本内部处理斜体
                    italic_parts = re.split(r'(\*[^*]+\*)', bold_content)

                    for italic_part in italic_parts:
                        if italic_part.startswith('*') and italic_part.endswith('*') and len(italic_part) >= 3:
                            # 加粗+斜体
                            run = paragraph.add_run(italic_part[1:-1])
                            run.bold = True
                            run.italic = True
                        else:
                            # 只加粗
                            if italic_part:
                                run = paragraph.add_run(italic_part)
                                run.bold = True
                else:
                    # 3. 处理斜体文本 *text* (低优先级)
                    italic_parts = re.split(r'(\*[^*]+\*)', bold_part)

                    for italic_part in italic_parts:
                        if italic_part.startswith('*') and italic_part.endswith('*') and len(italic_part) >= 3:
                            # 斜体文本
                            run = paragraph.add_run(italic_part[1:-1])
                            run.italic = True
                        else:
                            # 普通文本（包含emoji）
                            if italic_part:
                                paragraph.add_run(italic_part)

# Forensic Accounting Prompt Template
forensic_template = """{} ({})
Forensic and accounts investigation, As of today's date [{}]
Stock Price (Previous Close): {} {}
Market Cap: {} {}

Can you perform forensic accounting on the said company using their quarterly and annual reports?
Where it is possible to do so, please state definitively (a) whether accounting tricks most probably have been applied, (b) the probable accounting consequence (e.g. inflating sales, capitalizing cost so they don't deducted from income statements, (c ) hiding or deferring financial and repayment liabilities.
Please consider the extent of the most probable tricks and intent and conclude this analysis of whether there is high probability of extensive accounting problems.
For each section and sub-section, can you please state the section heading, and if nothing bad is found you will just say so. This allows users better understanding that you have covered everything.

Type - Revenue Recognition Tricks
These involve prematurely or fictitiously booking sales to inflate top-line growth.
Trick - Explanation
Channel Stuffing:
Flooding distributors with excess inventory at period-end to recognize revenue early, even without real demand. Example: Bristol-Myers Squibb inflated earnings by $1.5 billion through wholesaler overloads.
Bill-and-Hold Sales:
Booking revenue for goods "sold" but still held by the company for the buyer, often with side agreements allowing returns. Example: Alere Inc. recognized $24 million prematurely via third-party storage deals.
Round-Tripping
Circular transactions where funds are swapped between related parties (e.g., via loans disguised as sales) to create fake revenue. Example: Used in the Enron scandal to fabricate energy trading volumes.
Fake Sales (Fictitious Revenue)
Inventing sales from bogus invoices or ghost customers. Example: Luckin Coffee fabricated $300 million in 2019 revenue, leading to delisting.
Improper Timing of Revenue Recognition
Accelerating future sales into the current period or delaying them post-targets. Example: Marvell Technology pulled in 5-16% of quarterly revenue from future periods.
Third-Party Transactions
Deceptive deals like consignment sales masked as outright purchases. Often overlaps with bill-and-hold.

Type - Expense Manipulation
These defer or hide costs to overstate current profits.
Capitalizing Expenses (Improper Capitalization)
Treating operating costs (e.g., R&D or maintenance) as long-term assets to amortize over time instead of expensing immediately. Example: WorldCom capitalized $9 billion in line costs, inflating assets.
Cookie Jar Reserves
Over-provisioning for future expenses (e.g., excessive bad debt allowances) then reversing them in good periods to boost earnings. Example: Sunbeam used reserves to smooth income across quarters.
Deferred Expenses
Delaying recognition of costs through arbitrary accruals or inventory over-allocation. Often part of broader expense schemes.
Other Improper Expense Recognition
Understating liabilities like warranties or skipping impairments. Example: Celadon Group hid $20 million in asset write-downs via fake sales.

Type - Earnings Management
These smooth or manipulate overall profitability to meet targets or mislead analysts.
Trick - Explanation
Big Bath Accounting
Taking large one-time write-offs during bad periods to "clean the slate," making future earnings look stronger. Example: Used by banks post-2008 crisis to dump losses.
Earnings Smoothing
Averaging income over periods via reserves or estimates to avoid volatility. Example: General Electric smoothed earnings through insurance reserve manipulations.
Misclassifying Non-Recurrent Items
Labeling ongoing costs as "one-time" to exclude them from core earnings. Often tied to non-GAAP tweaks.
Fraudulent Management Estimates
Biasing subjective judgments (e.g., depreciation rates) to shift income. Example: Computer Sciences Corp. used flawed models, settling for $190 million.
Misleading Non-GAAP Reporting
Adjusting metrics like "pro forma" earnings to exclude real expenses. Example: Brixmor Property Group inflated same-store metrics, fined $7 million.

Type - Balance Sheet Shenanigans
These distort asset/liability portrayals to hide debt or overstate value.
Trick - Explanation
Off-Balance Sheet Financing
Hiding debt through special purpose entities (SPEs) or leases not consolidated. Example: Enron's SPEs concealed $13 billion in liabilities.
Improper Asset Valuation
Overvaluing inventory, intangibles, or investments via aggressive assumptions. Example: Valeant Pharmaceuticals inflated drug asset values pre-crash.
Understated Liabilities
Failing to accrue contingencies like lawsuits or pensions. Example: Toshiba understated $1.2 billion in cost overruns in 2015.

Type - Cash Flow Manipulation
These make operating cash look healthier than it is.
Trick - Explanation
Misclassifying Cash Flows
Shifting operating outflows to investing/financing (e.g., calling vendor prepayments "investments"). Example: Used by tech firms to boost free cash flow metrics.
Timing Tricks
Delaying supplier payments or accelerating collections to window-dress period-end cash. Often overlaps with revenue timing.

Type - Tax and Regulatory Evasion
These exploit loopholes to reduce reported taxes or skirt rules.
Trick - Explanation
Transfer Pricing Abuse
Shifting profits to low-tax jurisdictions via inflated inter-company sales. Example: Apple faced EU fines of €13 billion for Irish transfer pricing.
Tax Deferral Games
Aggressive use of NOL carryforwards or hybrid instruments to defer taxes indefinitely. Example: General Electric deferred billions through offshore structures.
Misleading Forecasts or Projections
Issuing false guidance to avoid regulatory scrutiny on shortfalls. Example: Walgreens reaffirmed optimistic merger projections, settling for $34.5 million.
Excessive use of supplier financing
Normal supplier financing is expressed in average days to payment. It can be between 30 days to 45 days. Some companies stretch that out to over 100 days or even 150 days, and they can bully suppliers because of their purchase size. This saves on the company having to borrow more to sustain its business. But in fact the financial conditions at this company is tighter and actual operational profits are much lower than stated in the accounts.

Cross-Cutting: Inadequate Internal Controls (ICFR)
This isn't a specific trick but enables many others by failing to detect/prevent them. Example: Monsanto's weak rebate controls led to an $80 million penalty.

This integrated list covers the most prevalent tactics from forensic accounting literature (e.g., SEC enforcement actions). It's more granular than my original but avoids redundancy.
Please add Traffic-light flag: 🟢, 🟡 , 🔴 to those sections where significant bad situations are found.
At the end of your report will you please compile a table showing the serious forensic problems found and a summary statement of effects on sales or profitability or on financial stability.
"""

########################## 主程序 ##########################

# Perplexity API 函数
def call_perplexity_api(prompt, max_retries=3):
    """调用 Perplexity Sonar Deep Research API"""
    headers = {
        "Authorization": f"Bearer {PERPLEXITY_API_KEY}",
        "Content-Type": "application/json"
    }

    data = {
        "model": "sonar-deep-research",
        "messages": [
            {
                "role": "system",
                "content": "You are a highly intelligent forensic accounting expert and AI assistant with deep research capabilities."
            },
            {
                "role": "user",
                "content": prompt
            }
        ]
    }

    for attempt in range(max_retries):
        try:
            response = requests.post(
                PERPLEXITY_API_URL,
                headers=headers,
                json=data,
                timeout=3600  # 1小时超时，适应深度研究模式
            )
            response.raise_for_status()
            result = response.json()

            if result and "choices" in result and len(result["choices"]) > 0:
                content = result["choices"][0]["message"]["content"]
                # 移除 <think> 标签内容（推理过程）
                content = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL)
                return content.strip()
            else:
                raise Exception("API返回格式异常")

        except Exception as e:
            print(f"⚠️  API调用失败 (尝试 {attempt + 1}/{max_retries}): {str(e)}")
            if attempt < max_retries - 1:
                time.sleep(5)
            else:
                raise

    return None


# =============================================================================
# Main Execution
# =============================================================================

# Validate configuration before starting
_validate_config()

# 读取Excel文件获取公司列表和股票代码
today = datetime.datetime.now().strftime("%Y-%m-%d")
today_short = datetime.datetime.now().strftime("%m%d")  # 获取月日格式，如0905
excel_filename = f"List - {today_short}.xlsx"  # 匹配实际文件名格式
# 使用脚本所在目录作为基础路径
script_dir = os.path.dirname(os.path.abspath(__file__))
excel_path = os.path.join(script_dir, excel_filename)

print(f"📊 正在读取Excel文件: {excel_path}")

try:
    # 读取Excel文件
    df = pd.read_excel(excel_path, skiprows=0)

    # 从A列读取公司名称，从B列读取股票代码，从F列读取文件名
    companies = df.iloc[:, 0].dropna().tolist()  # A列（第0列）
    tickers = df.iloc[:, 1].dropna().tolist()    # B列（第1列）
    file_names = df.iloc[:, 5].dropna().tolist() # F列（第5列）

    # 确保三个列表长度一致
    min_length = min(len(companies), len(tickers), len(file_names))
    companies = companies[:min_length]
    tickers = tickers[:min_length]
    file_names = file_names[:min_length]

    print(f"✅ 成功读取 {len(companies)} 家公司:")
    for i, (company, ticker, filename) in enumerate(zip(companies, tickers, file_names), 1):
        print(f"   {i}. {company} ({ticker}) - 文件名: {filename}")

except FileNotFoundError:
    print(f"❌ 错误: 找不到文件 {excel_path}")
    print(f"📝 请确保文件存在且命名格式为 'List-{today_short}.xlsx'")
    exit(1)
except Exception as e:
    print(f"❌ 读取Excel文件时发生错误: {e}")
    exit(1)

start_time = datetime.datetime.now()

# Create output directory with today's date（基于脚本所在目录）
output_base_dir = os.path.join(script_dir, "output")
output_dir = os.path.join(output_base_dir, today)
if not os.path.exists(output_dir):
    os.makedirs(output_dir, exist_ok=True)
    print(f"✅ Created output directory: {output_dir}")

# 初始化运行日志
run_log = []
successful_reports = []
failed_reports = []

print(f"\n🚀 开始处理 {len(companies)} 家公司的取证会计分析报告...")
print(f"📅 日期: {today}")
print("=" * 60)

run_log.append(f"🚀 开始处理 {len(companies)} 家公司的取证会计分析报告...")
run_log.append(f"📅 日期: {today}")
run_log.append("=" * 60)

# 初始化 Yahoo Finance 数据获取器
yahoo_fetcher = YahooFinanceDataFetcher()
yahoo_connected = yahoo_fetcher.connect()

if not yahoo_connected:
    print("⚠️ Yahoo Finance 连接失败，将使用空值作为股价和市值数据")

for i, (company, ticker, file_name) in enumerate(zip(companies, tickers, file_names), 1):
    company_start_time = datetime.datetime.now()
    print(f"\n📊 正在处理第 {i}/{len(companies)} 家公司: {company} ({ticker}) - 文件名: {file_name}")
    print(f"⏰ 开始时间: {company_start_time.strftime('%H:%M:%S')}")

    run_log.append(f"\n📊 正在处理第 {i}/{len(companies)} 家公司: {company} ({ticker}) - 文件名: {file_name}")
    run_log.append(f"⏰ 开始时间: {company_start_time.strftime('%H:%M:%S')}")

    # 获取股价和市值数据
    stock_data = None
    stock_price_text = "N/A"
    market_cap_text = "N/A"
    currency = ""

    if yahoo_connected:
        try:
            print(f"💰 从 Yahoo Finance 获取 {ticker} 的数据...")
            stock_data = yahoo_fetcher.get_stock_data(ticker)

            if stock_data and stock_data['last_price'] is not None:
                stock_price_text = f"{stock_data['last_price']:.2f}"
                currency = stock_data.get('currency', '')

                if stock_data['market_cap'] is not None:
                    # 格式化市值
                    market_cap_value = stock_data['market_cap']
                    if market_cap_value >= 1e12:
                        market_cap_text = f"{market_cap_value/1e12:.2f}T"
                    elif market_cap_value >= 1e9:
                        market_cap_text = f"{market_cap_value/1e9:.2f}B"
                    elif market_cap_value >= 1e6:
                        market_cap_text = f"{market_cap_value/1e6:.2f}M"
                    else:
                        market_cap_text = f"{market_cap_value:,.0f}"

                print(f"✅ 获取成功: 前收盘价={stock_price_text} {currency}, 市值={market_cap_text} {currency}")
                run_log.append(f"✅ 股价数据: 前收盘价={stock_price_text} {currency}, 市值={market_cap_text} {currency}")
            else:
                print(f"⚠️ 未能获取 {ticker} 的有效数据")
                run_log.append(f"⚠️ 未能获取 {ticker} 的有效数据")

        except Exception as e:
            print(f"❌ 获取股价数据时出错: {e}")
            run_log.append(f"❌ 获取股价数据时出错: {e}")

    # Generate the prompt with stock price and market cap data
    prompt = forensic_template.format(company, ticker, today, stock_price_text, currency, market_cap_text, currency)

    # 重试机制
    max_retries = 3
    retry_count = 0
    success = False

    while retry_count < max_retries and not success:
        try:
            print(f"🔄 尝试第 {retry_count + 1}/{max_retries} 次调用AI API...")
            run_log.append(f"🔄 尝试第 {retry_count + 1}/{max_retries} 次调用AI API...")

            # 调用 Perplexity API
            print("⏳ 等待Perplexity AI生成取证会计分析报告...")
            run_log.append("⏳ 等待Perplexity AI生成取证会计分析报告...")

            response_content = call_perplexity_api(prompt, max_retries=1)

            # 验证响应内容
            if response_content and len(response_content.strip()) > 100:
                # 转换Markdown为Word格式
                print("📝 转换Markdown格式到Word...")
                run_log.append("📝 转换Markdown格式到Word...")

                # Create a new Word document
                doc = Document()
                doc = convert_markdown_to_word(response_content, doc)

                # Save file to the dated output directory with "FA -" prefix
                output_filename = os.path.join(output_dir, f"FA - {file_name}.docx")
                doc.save(output_filename)

                company_end_time = datetime.datetime.now()
                processing_time = (company_end_time - company_start_time).total_seconds()

                print(f"✅ 成功保存: {output_filename}")
                print(f"📄 内容长度: {len(response_content)} 字符")
                print(f"📝 已转换为格式化Word文档")
                print(f"⏱️  处理耗时: {processing_time:.1f} 秒")

                run_log.append(f"✅ 成功保存: {output_filename}")
                run_log.append(f"📄 内容长度: {len(response_content)} 字符")
                run_log.append(f"📝 已转换为格式化Word文档")
                run_log.append(f"⏱️  处理耗时: {processing_time:.1f} 秒")

                successful_reports.append({
                    'company': company,
                    'ticker': ticker,
                    'filename': output_filename,
                    'content_length': len(response_content),
                    'processing_time': processing_time,
                    'completed_at': company_end_time.strftime('%H:%M:%S')
                })

                success = True

            else:
                print(f"⚠️  AI返回内容为空或过短，准备重试...")
                run_log.append(f"⚠️  AI返回内容为空或过短，准备重试...")
                retry_count += 1
                if retry_count < max_retries:
                    wait_time = 10 * retry_count
                    print(f"⏸️  等待 {wait_time} 秒后重试...")
                    run_log.append(f"⏸️  等待 {wait_time} 秒后重试...")
                    time.sleep(wait_time)

        except Exception as e:
            print(f"❌ 错误: {e}")
            run_log.append(f"❌ 错误: {e}")
            retry_count += 1
            if retry_count < max_retries:
                wait_time = 15 * retry_count
                print(f"⏸️  等待 {wait_time} 秒后重试...")
                run_log.append(f"⏸️  等待 {wait_time} 秒后重试...")
                time.sleep(wait_time)

    if not success:
        print(f"🚫 {company} 处理失败，已达到最大重试次数")
        run_log.append(f"🚫 {company} 处理失败，已达到最大重试次数")

        # 创建一个错误文件记录
        error_filename = os.path.join(output_dir, f"ERROR - FA - {file_name}.txt")
        with open(error_filename, 'w', encoding='utf-8') as f:
            f.write(f"处理失败\n公司: {company}\n股票代码: {ticker}\n文件名: {file_name}\n时间: {datetime.datetime.now()}\n")
        print(f"📝 错误记录已保存: {error_filename}")
        run_log.append(f"📝 错误记录已保存: {error_filename}")

        failed_reports.append({
            'company': company,
            'ticker': ticker,
            'error_file': error_filename,
            'failed_at': datetime.datetime.now().strftime('%H:%M:%S')
        })

    # 公司间等待时间
    if i < len(companies):
        print(f"⏸️  等待 5 秒后处理下一家公司...")
        run_log.append(f"⏸️  等待 5 秒后处理下一家公司...")
        time.sleep(5)

# 计算总耗时
end_time = datetime.datetime.now()
total_time = (end_time - start_time).total_seconds()

print(f"\n🎉 所有公司处理完成!")
print(f"📁 文件保存位置: {output_dir}")
print(f"⏰ 完成时间: {end_time.strftime('%H:%M:%S')}")
print(f"⏱️  总耗时: {total_time:.1f} 秒 ({total_time/60:.1f} 分钟)")

run_log.append(f"\n🎉 所有公司处理完成!")
run_log.append(f"📁 文件保存位置: {output_dir}")
run_log.append(f"⏰ 完成时间: {end_time.strftime('%H:%M:%S')}")
run_log.append(f"⏱️  总耗时: {total_time:.1f} 秒 ({total_time/60:.1f} 分钟)")

# 生成运行报告Word文档
report_doc = Document()

# 添加标题
title = report_doc.add_heading('Forensic Accounting Analysis - Execution Report', 0)
title.alignment = 1  # 居中

# 添加总结段落
summary = report_doc.add_heading('📊 Execution Summary', level=1)

# 创建总结表格
summary_table = report_doc.add_table(rows=1, cols=2)
summary_table.style = 'Table Grid'

# 表格标题行
hdr_cells = summary_table.rows[0].cells
hdr_cells[0].text = 'Item'
hdr_cells[1].text = 'Result'

# 添加总结数据
summary_data = [
    ('Execution Date', today),
    ('Start Time', start_time.strftime('%H:%M:%S')),
    ('End Time', end_time.strftime('%H:%M:%S')),
    ('Total Duration', f"{total_time:.1f} seconds ({total_time/60:.1f} minutes)"),
    ('Total Companies to Process', str(len(companies))),
    ('Successful Reports Generated', str(len(successful_reports))),
    ('Failed Reports', str(len(failed_reports))),
    ('Success Rate', f"{len(successful_reports)/len(companies)*100:.1f}%")
]

for item, value in summary_data:
    row_cells = summary_table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = value

# 添加成功报告详情
if successful_reports:
    success_heading = report_doc.add_heading('✅ Successfully Generated Reports', level=1)

    success_table = report_doc.add_table(rows=1, cols=5)
    success_table.style = 'Table Grid'

    hdr_cells = success_table.rows[0].cells
    hdr_cells[0].text = 'Company Name'
    hdr_cells[1].text = 'Stock Code'
    hdr_cells[2].text = 'Completion Time'
    hdr_cells[3].text = 'Processing Time (sec)'
    hdr_cells[4].text = 'Content Length (chars)'

    for report in successful_reports:
        row_cells = success_table.add_row().cells
        row_cells[0].text = report['company']
        row_cells[1].text = report['ticker']
        row_cells[2].text = report['completed_at']
        row_cells[3].text = f"{report['processing_time']:.1f}"
        row_cells[4].text = str(report['content_length'])

# 添加失败报告详情
if failed_reports:
    failed_heading = report_doc.add_heading('❌ Failed Reports', level=1)

    failed_table = report_doc.add_table(rows=1, cols=3)
    failed_table.style = 'Table Grid'

    hdr_cells = failed_table.rows[0].cells
    hdr_cells[0].text = 'Company Name'
    hdr_cells[1].text = 'Stock Code'
    hdr_cells[2].text = 'Failed Time'

    for report in failed_reports:
        row_cells = failed_table.add_row().cells
        row_cells[0].text = report['company']
        row_cells[1].text = report['ticker']
        row_cells[2].text = report['failed_at']

# 添加详细运行日志
log_heading = report_doc.add_heading('📋 Detailed Execution Log', level=1)

for log_entry in run_log:
    report_doc.add_paragraph(log_entry)

# 保存运行报告
report_filename = os.path.join(output_dir, f"FA Execution Report - {today} - {end_time.strftime('%H%M%S')}.docx")
report_doc.save(report_filename)

# 断开 Yahoo Finance 连接
if yahoo_connected:
    yahoo_fetcher.disconnect()

print(f"📄 运行报告已保存: {report_filename}")
print(f"📈 成功率: {len(successful_reports)}/{len(companies)} ({len(successful_reports)/len(companies)*100:.1f}%)")
