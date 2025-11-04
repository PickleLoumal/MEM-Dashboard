#pip install python-docx
#pip install xai-sdk
#pip install lseg-data
import os
import datetime
import time
import re
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# XAI SDK 配置
os.environ["XAI_API_KEY"] = "xai-cvJvJSW2ZLXqypPjOHmDFhxOCsAL6y0I5xUS05Z2l1C3b1U9UeGNeJMi8NLWYG5j17DROUiRMOgXuaYZ"
from xai_sdk import Client
from xai_sdk.chat import user, system

# LSEG Data API 配置
import lseg.data as ld

# 设置API密钥环境变量
def setup_lseg_api_config(api_key):
    """
    设置LSEG API配置
    """
    try:
        os.environ['RDP_APP_KEY'] = api_key
        print("✅ LSEG API配置设置成功")
        return True
    except Exception as e:
        print(f"⚠️ LSEG API配置设置失败: {e}")
        return False

class LSEGDataFetcher:
    """
    LSEG数据获取器，专门用于获取股价和市值数据
    """
    
    def __init__(self, api_key='4a63f4c1e76243efb869f5a8a905a001ead21a42'):
        self.api_key = api_key
        self.session = None
        setup_lseg_api_config(self.api_key)
        
    def connect(self):
        """
        建立LSEG连接
        """
        try:
            print("🔌 连接到LSEG数据平台...")
            
            # 尝试多种连接方式
            try:
                self.session = ld.open_session()
                print("✅ 成功连接到LSEG (默认)")
                return True
            except:
                try:
                    self.session = ld.open_session('desktop.workspace')
                    print("✅ 成功连接到LSEG (Desktop)")
                    return True
                except:
                    try:
                        self.session = ld.open_session('deployed')
                        print("✅ 成功连接到LSEG (Deployed)")
                        return True
                    except Exception as e:
                        print(f"❌ 所有LSEG连接方式都失败: {e}")
                        return False
                        
        except Exception as e:
            print(f"❌ LSEG连接失败: {e}")
            return False
    
    def get_stock_data(self, symbol):
        """
        获取股票的前一天收盘价和市值数据
        
        参数:
        symbol: 股票代码 (如 '0700.HK', 'AAPL.O')
        
        返回:
        dict: {'last_price': 前一天收盘价, 'market_cap': 市值, 'currency': 货币}
        """
        try:
            if not self.session:
                print("请先建立LSEG连接")
                return None
                
            print(f"📊 获取 {symbol} 的前一天收盘价和市值数据...")
            
            # 获取股价和市值字段 - 使用前一天收盘价 HST_CLOSE
            fields = [
                'HST_CLOSE',                     # 前一天收盘价 ✅
                'TR.CLOSEPRICE',                 # 收盘价 (备用) ✅
                'TR.PriceClose',                 # 收盘价 (备用2) ✅
                'CF_CLOSE',                      # 收盘价 (通用字段) ✅
                'TR.CompanyMarketCap',           # 市值 ✅
                'TR.CompanyMarketCapitalization', # 备用市值 ✅
                'TR.Currency',                   # 交易货币 (主要字段) 💰
                'CF_CURRENCY',                   # 货币字段 (备用) 💰
                'TR.TradingItemCurrency',        # 交易货币 (详细) 💰
                'TR.ExchangeCurrency',           # 交易所货币 💰
                'TR.PriceCurrency',              # 价格货币 💰
                'CURRENCY',                      # 简单货币字段 💰
                'CCY'                            # 货币简写 💰
            ]
            
            data = ld.get_data(
                universe=[symbol],
                fields=fields
            )
            
            if data is not None and not data.empty:
                # 调试：显示实际返回的列名
                print(f"🔍 调试：实际返回的列名: {list(data.columns)}")
                
                # 获取实际的列名映射
                price_value = None
                market_cap_value = None
                currency_value = None
                
                # 获取前一天收盘价 - 尝试多个可能的列名
                if 'HST_CLOSE' in data.columns:
                    price_value = data['HST_CLOSE'].iloc[0] if not data['HST_CLOSE'].isna().iloc[0] else None
                elif 'TR.CLOSEPRICE' in data.columns:
                    price_value = data['TR.CLOSEPRICE'].iloc[0] if not data['TR.CLOSEPRICE'].isna().iloc[0] else None
                elif 'Close Price' in data.columns:
                    price_value = data['Close Price'].iloc[0] if not data['Close Price'].isna().iloc[0] else None
                elif 'TR.PriceClose' in data.columns:
                    price_value = data['TR.PriceClose'].iloc[0] if not data['TR.PriceClose'].isna().iloc[0] else None
                elif 'CF_CLOSE' in data.columns:
                    price_value = data['CF_CLOSE'].iloc[0] if not data['CF_CLOSE'].isna().iloc[0] else None
                
                # 获取市值 - 使用实际返回的列名
                if 'Company Market Cap' in data.columns:
                    market_cap_value = data['Company Market Cap'].iloc[0] if not data['Company Market Cap'].isna().iloc[0] else None
                elif 'Company Market Capitalization' in data.columns:
                    market_cap_value = data['Company Market Capitalization'].iloc[0] if not data['Company Market Capitalization'].isna().iloc[0] else None
                
                # 获取货币 - 尝试多个可能的列名（按优先级排序）
                currency_fields = ['TR.Currency', 'CF_CURRENCY', 'TR.TradingItemCurrency', 'TR.ExchangeCurrency', 'TR.PriceCurrency', 'CURRENCY', 'CCY']
                for curr_field in currency_fields:
                    if curr_field in data.columns:
                        temp_currency = data[curr_field].iloc[0] if not data[curr_field].isna().iloc[0] else None
                        if temp_currency is not None and str(temp_currency).upper() not in ['NAN', 'NULL', 'NONE', '<NA>']:
                            currency_value = temp_currency
                            break
                
                result = {
                    'symbol': symbol,
                    'last_price': price_value,
                    'market_cap': market_cap_value,
                    'currency': currency_value
                }
                
                print(f"✅ 获取成功: 前收盘价={result['last_price']}, 市值={result['market_cap']}, 货币={result['currency']}")
                return result
            else:
                print(f"⚠️ 未获取到 {symbol} 的数据，尝试备用字段...")
                
                # 尝试备用字段
                backup_fields = [
                    'LAST',              # 最新价格 (备用)
                    'TR.MarketCap',      # 市值 (备用)  
                    'TR.CompanyMarketCap', # 市值 (备用2)
                    'DSPLY_NAME'         # 显示名称
                ]
                
                backup_data = ld.get_data(
                    universe=[symbol],
                    fields=backup_fields
                )
                
                if backup_data is not None and not backup_data.empty:
                    # 使用同样的方式获取备用数据
                    backup_price = None
                    backup_market_cap = None
                    
                    if 'LAST' in backup_data.columns:
                        backup_price = backup_data['LAST'].iloc[0] if not backup_data['LAST'].isna().iloc[0] else None
                    
                    # 尝试多个市值字段
                    for cap_field in ['TR.MarketCap', 'Market Cap', 'Company Market Cap', 'TR.CompanyMarketCap']:
                        if cap_field in backup_data.columns:
                            backup_market_cap = backup_data[cap_field].iloc[0] if not backup_data[cap_field].isna().iloc[0] else None
                            if backup_market_cap is not None:
                                break
                    
                    result = {
                        'symbol': symbol,
                        'last_price': backup_price,
                        'market_cap': backup_market_cap,
                        'currency': 'CNY'  # 中国股票默认货币
                    }
                    
                    print(f"✅ 备用方法获取成功: 前收盘价={result['last_price']}, 市值={result['market_cap']}")
                    return result
                else:
                    print(f"❌ 无法获取 {symbol} 的数据")
                    return None
                    
        except Exception as e:
            print(f"❌ 获取 {symbol} 前收盘价数据时出错: {e}")
            return None
    
    def disconnect(self):
        """
        断开LSEG连接
        """
        try:
            if self.session:
                ld.close_session()
                print("✅ LSEG连接已断开")
        except Exception as e:
            print(f"❌ 断开LSEG连接失败: {e}")

# XAI客户端初始化
client = Client(
    api_key=os.getenv("XAI_API_KEY"),
    timeout=3600,  # Override default timeout with longer timeout for reasoning models
)

def convert_markdown_to_word(markdown_text, doc):
    """
    将Markdown文本转换为Word文档格式，严格按照markdown语法标准
    """
    lines = markdown_text.split('\n')
    in_code_block = False
    code_language = ""
    
    for line in lines:
        # 处理代码块
        if line.strip().startswith('```'):
            if not in_code_block:
                # 开始代码块
                in_code_block = True
                code_language = line.strip()[3:].strip()  # 获取语言标识
                continue
            else:
                # 结束代码块
                in_code_block = False
                code_language = ""
                continue
                
        if in_code_block:
            # 代码块内容用等宽字体，保持原始格式
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            continue
        
        # 跳过空行但保留段落间距
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # 处理各级标题 - 必须是行首且后面有空格
        if line.startswith('# ') and not line.startswith('## '):
            heading = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## ') and not line.startswith('### '):
            heading = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### ') and not line.startswith('#### '):
            heading = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### ') and not line.startswith('##### '):
            heading = doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('##### ') and not line.startswith('###### '):
            heading = doc.add_heading(line[6:].strip(), level=5)
        
        # 处理无序列表 - 支持缩进
        elif re.match(r'^[ \t]*[-*+] ', line):
            # 计算缩进级别
            indent_match = re.match(r'^([ \t]*)', line)
            indent_level = len(indent_match.group(1).expandtabs(4)) // 4 if indent_match else 0
            
            # 提取列表内容
            bullet_text = re.sub(r'^[ \t]*[-*+] ', '', line)
            
            # 创建列表项
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, bullet_text)
            
        # 处理有序列表
        elif re.match(r'^[ \t]*\d+\. ', line):
            # 提取列表内容
            bullet_text = re.sub(r'^[ \t]*\d+\. ', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, bullet_text)
            
        # 处理水平分割线
        elif line.strip() in ['---', '***', '___'] or re.match(r'^[ \t]*[-*_]{3,}[ \t]*$', line):
            # 添加水平线
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run('─' * 50)
            run.font.color.rgb = RGBColor(128, 128, 128)
            
        # 处理引用块 (以 > 开头)
        elif line.strip().startswith('> '):
            quote_text = line.strip()[2:]
            p = doc.add_paragraph()
            p.left_indent = Inches(0.5)
            p.paragraph_format.left_indent = Inches(0.5)
            add_formatted_text(p, quote_text)
            
        # 处理普通段落
        else:
            if line.strip():
                p = doc.add_paragraph()
                add_formatted_text(p, line.strip())
    
    return doc

def add_formatted_text(paragraph, text):
    """
    向段落添加格式化文本，按照正确的markdown优先级处理
    """
    # 按照markdown标准的优先级处理：先处理加粗，再处理斜体，最后处理行内代码
    
    # 1. 先处理行内代码 `code` (最高优先级，不能被其他格式影响)
    code_parts = re.split(r'(`[^`]*`)', text)
    
    for code_part in code_parts:
        if code_part.startswith('`') and code_part.endswith('`') and len(code_part) >= 2:
            # 行内代码 - 直接添加，不处理其他格式
            run = paragraph.add_run(code_part[1:-1])
            run.font.name = 'Courier New'
        else:
            # 2. 处理加粗文本 **text** (高优先级)
            bold_parts = re.split(r'(\*\*[^*]*\*\*)', code_part)
            
            for bold_part in bold_parts:
                if bold_part.startswith('**') and bold_part.endswith('**') and len(bold_part) >= 4:
                    # 加粗文本 - 内容还可能包含斜体
                    bold_content = bold_part[2:-2]
                    # 在加粗文本内部处理斜体
                    italic_parts = re.split(r'(\*[^*]+\*)', bold_content)
                    
                    for italic_part in italic_parts:
                        if italic_part.startswith('*') and italic_part.endswith('*') and len(italic_part) >= 3:
                            # 加粗+斜体
                            run = paragraph.add_run(italic_part[1:-1])
                            run.bold = True
                            run.italic = True
                        else:
                            # 只加粗
                            if italic_part:
                                run = paragraph.add_run(italic_part)
                                run.bold = True
                else:
                    # 3. 处理斜体文本 *text* (低优先级)
                    italic_parts = re.split(r'(\*[^*]+\*)', bold_part)
                    
                    for italic_part in italic_parts:
                        if italic_part.startswith('*') and italic_part.endswith('*') and len(italic_part) >= 3:
                            # 斜体文本
                            run = paragraph.add_run(italic_part[1:-1])
                            run.italic = True
                        else:
                            # 普通文本
                            if italic_part:
                                paragraph.add_run(italic_part)



########################## 主程序 ##########################

# 读取Excel文件获取公司列表和股票代码
today = datetime.datetime.now().strftime("%Y-%m-%d")
today_short = datetime.datetime.now().strftime("%m%d")  # 获取月日格式，如0905
excel_filename = f"List - {today_short}.xlsx"  # 匹配实际文件名格式
excel_path = os.path.join(r"C:\Users\User\Desktop\automate", excel_filename)

print(f"📊 正在读取Excel文件: {excel_path}")

try:
    # 读取Excel文件，从第二行开始（skiprows=0表示不跳过任何行，从第一行开始读取）
    # 如果第一行是标题，pandas会自动识别；如果要跳过标题行，可以用skiprows=1
    df = pd.read_excel(excel_path, skiprows=0)
    
    # 从A列读取公司名称，从B列读取股票代码，从F列读取文件名
    companies = df.iloc[:, 0].dropna().tolist()  # A列（第0列）
    HKShare = df.iloc[:, 1].dropna().tolist()    # B列（第1列）
    file_names = df.iloc[:, 5].dropna().tolist() # F列（第5列）
    
    # 确保三个列表长度一致
    min_length = min(len(companies), len(HKShare), len(file_names))
    companies = companies[:min_length]
    HKShare = HKShare[:min_length]
    file_names = file_names[:min_length]
    
    print(f"✅ 成功读取 {len(companies)} 家公司:")
    for i, (company, ticker, filename) in enumerate(zip(companies, HKShare, file_names), 1):
        print(f"   {i}. {company} ({ticker}) - 文件名: {filename}")
    
except FileNotFoundError:
    print(f"❌ 错误: 找不到文件 {excel_path}")
    print(f"📝 请确保文件存在且命名格式为 'list-{today_short}.xlsx'")
    exit(1)
except Exception as e:
    print(f"❌ 读取Excel文件时发生错误: {e}")
    exit(1)
template = """{}\n{}\nAs of today [{}]\nStock Price (Previous Close): {} {}\nMarket Cap: {} {}\nObjective: Create a concise investment summary (maximum 5 pages, ~450-600 words) for the said company (including its major business divisions and segments, subsidiaries and parent companies), covering business overview, financial stability, valuation, anomalies, customer segments, demand trends, competitive landscape, and a Buy/Hold/Sell recommendation.. Use regulatory filings (10-K, 10-Q), investor transcripts, industry reports (e.g., McKinsey, Deloitte, EY), analyst insights, and market data, updated to the current date). Include a section on financial stability and debt levels, and highlight leading investment firms’ views with target prices. Focus on the most recommended Buy/Hold/Sell stance with pros and cons, avoiding analysis of all three options. Ensure the response is not professional investment advice.\nPlease use bullet points whenever possible.\nPrompt: Please generate an investment summary for the said Company as of the current date stated above, adhering to the following structure and requirements:\nSources of information and data: Please make sure that among wide range of information and sources, you must include authoritative sources, including the :\nInformation published or available from the Company and its subsidiaries - including company websites and presentations, annual reports and quarterly reports (e.g. in the USA, forms 10K and 10Q, SEC filings and EDGAR database contents, but may be differently designated in other countries)\nThe MD&A and similar business and outlook statements, including opportunities and risks, and industry trends and changes.\nTranscripts (if available) of earnings call, investors conferences, and other statements of the company.\nRegulatory bodies publications and statistics and statement about the Company or, in general, about the industr(ies) it is in.\nIndustry specific ratios, and the industry median vs those as applied to the Company.\nBut please note that you are free and indeed encouraged to look much wider afield.\nIn respect of this chat - Please provide a link to each information or data source you used.\nConfirm you have used and not skipped over any of the said authoritative sources.\nSome limits - A current ratio below 1.3 is not considered financially healthy as concerns liquidity, unless the company is a cash business company, like McDonald’s or JD.com.\nTry to get as updated data and information as possible. Things can change day by day on the stock market.\nHeader:\nTitle: "Investment Summary", then add the name of the said Company\nInclude: Current date, stock price (close from the previous trading day), market cap, and recommended action (Buy, Hold, or Sell). Also add the name of the industry it is in. If more than one industry then name each of the industries it is in.\nBusiness Overview (1 paragraph):\nSummarize the company’s operations, major divisions, products/services, and key financials (e.g., recent FY sales, operating income, margins).\nInclude a 2 sentence explanation of the use of each of its products to its major customer segments.\nHighlight strengths (e.g., Technology, functionality and capabilities of products, brand equity, operational efficiencies) and challenges (e.g., market pressures, risks).\nUse FY data or year-to-date if more recent; specify fiscal year-end.\nFor each major division or product line / service: include its sales as a percentage of total group sales, and (if available), the gross profit margin and as a percentage of group profits.\nBusiness Performance \nBullet Points - \n(a) Sales growth in the past 5 years. Forecast for next year \n(b) Profit growth in the past 5 years. Forecast for next year. \n(c) operating cash flow increase\n(d) market share and ranking in its industry\nIndustry context -\nFor the industry in which this company is in (and if more than one, deal with each), please find out the following - \n(a) Product cycle maturity. \n(b) market size and growth rate CAGR\n(c) this company's market share and ranking \n(d) average sales growth over the past 3 years of this company compared to industry average\n(e) average EPS growth over the past 3 years of this company compared to industry average.\n(f) debt-to-total assets ratio of this company compared to industry average. \n(g) Is the industry cycle in an expansion phase or slowing down phase? (an example to illustrate this - in the insurance industry is it in a soft market or a hard market) \n(h) industry specific metrics (not the usual PE ratio or financial ratios. For example, in the offshore drilling industry, such a ratio include "rig count", and in the container shipping industry, the "World Container Index") and for each of them how does this company compare to the industry average?\nFinancial Stability and Debt Levels (1 paragraph):\nAssess financial stability, focusing on operating cash flow, dividend coverage, capex, and liquidity (e.g., cash on hand, current ratio).\nEvaluate debt levels (e.g., total debt, debt-to-equity ratio, debt-to-total assets ratio, interest coverage, Altman Z Score) relative to industry norms.\nHighlight any financial problems (e.g., high leverage, weak cash flow) or confirm prudent debt management.\nKey Financials and Valuation (bullet points):\nSales and Profitability: including (a) Most recent past financial year sale and Year on Year change compared to prior year, and forecast, (b) Sales and profitability performance of each division of operations;, (c ) operating profit margin and trends for the group,. Include forward perfornace guidance (sales, EPS, YoY change)\nValuation Metrics: P/E (TTM, vs. industry and historical average), PEG, dividend yield, and stock’s position in 52-week range.\nFinancial stability and debt levels: commonly used financial analysis and debt ratios and highlighting risks.\nIndustry Specific metrics: Those which are specific and considered important to the industry that this company is in, as applied to this company and as compared to industry averages. Find 3 industry specific metrics for the industry the said Company is in, and also find those same ratios for the said company, and compare them and rate how the Company compares. Provide brief observations about the comparison and what it means for the said Company.\n(Examples of industry specific metrics : Banking industry - Capital adequacy ratio, loan-to-deposit ratio, NIM, NII. Example in the Semiconductor wafer fab industry - Book-to-bill ratio, die size as in 2nm, Yield. Example in the Airline industry: Loading factor, breakeven loading factor. Example in the Offshore oil drilling services industry: Rig Count, day rate of rig rental.\nBig Trends and Big Events \nThose that are affecting the Industry that the Company is in (separately for each different business segments.\nFor each big trend or big events, explain the possible effects on companies in this industry in general and this Company specifically.\nCustomer Segments and Demand Trends (bullet points):\nMajor Segments: List top customer segments by sales (currency and %), e.g., Retail (supermarkets, mass merchandisers), Foodservice, International.\nForecast: Sales growth projections (next 2-3 years) per segment, with key drivers (e.g., innovation, market trends).\nCriticisms and Substitutes: Customer complaints (e.g., price, functionality) and substitute products with switching speed.\nCompetitive Landscape (bullet points):\nIndustry Dynamics: Assess competitiveness using concentration (CR4), margins, capacity utilization, growth (CAGR), and industry cycle stage.\nKey Competitors: List major competitors, estimated market shares, and operating margins.\nMoats: Identify sustainable moats (e.g. Technology, network effects, switching costs, economy of scale, hard to obtain or exclusive government permits or licenses, supply chain integration upstream or downstream, cost leadership, brands, distribution) and assess the company’s moats vs. competitors.\nKey battle fronts in this industry, e.g. Scale of operation, ownership of supply chains and sales channels, capex resources, technology, regulatory licensing, brand names and customer loyalty, customer switching costs, others? Pick the top one and explain how this company measures up to competitors.\nRisks and anomalies (bullet points):\nHighlight unusual findings (e.g., divisional sales drops vs. stable group profits, litigation costs, market volatility).\nExplain concerns and potential resolution (e.g., operational fixes, legal settlements).\nForecast and outlook (bullet points):\nInclude: Forecast sales and profits by management of the company, key growth (or decline) from specific product lines / service lines. Key reasons for growth / decline. Recent earnings surprise and reasons.\nLeading Investment Firms and Views (bullet points):\nList top firms (e.g., Piper Sandler, Goldman Sachs) and analysts covering the company/industry, with their latest ratings and target prices (include % upside/downside).\nInclude consensus rating and average target price (with range).\nRecommended Action: [Buy/Hold/Sell] (1 section):\nSelect the most appropriate recommendation (Buy, Hold, or Sell) based on analysis.\nExplain the reasons for the most appropriated recommendation, by stating the reasons for Pros (e.g., financial stability, growth potential, analyst optimism) and Cons (e.g., valuation risks, competitive pressures) in bullet points.\nIndustry Ratio and Metric Analysis:\nWhat are the important industry specific ratios and metrics in the industry that this company is in? Can you (a) give me the ratios and metrics for each one for this company, (b) comparison to industry average, (c) Trends for the industry and for the company?\nKey Takeaways (1 paragraph for each of the key takeaway points.):\nSummarize the company’s position, strengths, risks, and recommendation rationale.\nHighlight monitorable factors (e.g., innovation, risk resolution) for future opportunities.\nHave we missed out on some key or important points which would provide to us much better and proper understanding of this company and its business?\nAdditional Instructions:\nKeep the summary concise (~450-500 words, fitting 3 pages).\nUse data from regulatory filings, investor transcripts, industry reports (e.g., McKinsey, Deloitte, EY), and analyst notes (e.g., Piper Sandler, Goldman Sachs).\nCite sources at the end (e.g., 10-K/10-Q, industry reports, market data like Yahoo Finance).\nUse bullet points for clarity in sections 4-9.\nEnsure financial stability section explicitly addresses cash flow, debt, and any financial concerns.\nAvoid analyzing all three Buy/Hold/Sell options; focus on the most recommended with pros and cons.\nFormat as markdown within tag, with a unique artifact_id, consistent title (e.g., "[Company Name] Investment Summary.md"), and contentType="text/markdown"""
start_time = datetime.datetime.now()

# Create output directory with today's date
output_dir = os.path.join(r"C:\Users\User\Desktop\automate\output", today)
if not os.path.exists(output_dir):
    os.makedirs(output_dir)
    print(f"✅ Created output directory: {output_dir}")

# 初始化运行日志
run_log = []
successful_reports = []
failed_reports = []

print(f"\n🚀 开始处理 {len(companies)} 家公司的投资分析报告...")
print(f"📅 日期: {today}")
print("=" * 60)

run_log.append(f"🚀 开始处理 {len(companies)} 家公司的投资分析报告...")
run_log.append(f"📅 日期: {today}")
run_log.append("=" * 60)

# 初始化LSEG数据获取器
lseg_fetcher = LSEGDataFetcher()
lseg_connected = lseg_fetcher.connect()

if not lseg_connected:
    print("⚠️ LSEG连接失败，将使用空值作为股价和市值数据")

for i, (company, HKticker, file_name) in enumerate(zip(companies, HKShare, file_names), 1):
    company_start_time = datetime.datetime.now()
    print(f"\n📊 正在处理第 {i}/{len(companies)} 家公司: {company} ({HKticker}) - 文件名: {file_name}")
    print(f"⏰ 开始时间: {company_start_time.strftime('%H:%M:%S')}")
    
    run_log.append(f"\n📊 正在处理第 {i}/{len(companies)} 家公司: {company} ({HKticker}) - 文件名: {file_name}")
    run_log.append(f"⏰ 开始时间: {company_start_time.strftime('%H:%M:%S')}")
    
    # 获取股价和市值数据
    stock_data = None
    stock_price_text = "N/A"
    market_cap_text = "N/A"
    currency = ""
    
    if lseg_connected:
        try:
            print(f"💰 获取 {HKticker} 的前一天收盘价和市值...")
            stock_data = lseg_fetcher.get_stock_data(HKticker)
            
            if stock_data and stock_data['last_price'] is not None:
                stock_price_text = f"{stock_data['last_price']:.2f}"
                currency = stock_data.get('currency', '')
                
                if stock_data['market_cap'] is not None:
                    # 格式化市值 (根据实际数值大小格式化)
                    market_cap_value = stock_data['market_cap']
                    if market_cap_value >= 1e12:  # 万亿
                        market_cap_text = f"{market_cap_value/1e12:.2f}T"  # 万亿
                    elif market_cap_value >= 1e9:  # 十亿
                        market_cap_text = f"{market_cap_value/1e9:.2f}B"  # 十亿
                    elif market_cap_value >= 1e6:  # 百万
                        market_cap_text = f"{market_cap_value/1e6:.2f}M"  # 百万
                    else:
                        market_cap_text = f"{market_cap_value:,.0f}"  # 小于百万直接显示
                        
                print(f"✅ 获取成功: 前收盘价={stock_price_text} {currency}, 市值={market_cap_text} {currency}")
                run_log.append(f"✅ 股价数据: 前收盘价={stock_price_text} {currency}, 市值={market_cap_text} {currency}")
            else:
                print(f"⚠️ 未能获取 {HKticker} 的有效数据")
                run_log.append(f"⚠️ 未能获取 {HKticker} 的有效数据")
                
        except Exception as e:
            print(f"❌ 获取股价数据时出错: {e}")
            run_log.append(f"❌ 获取股价数据时出错: {e}")
    
    # Generate the prompt with stock price and market cap data
    prompt = template.format(company, HKticker, today, stock_price_text, currency, market_cap_text, currency)

    # 重试机制
    max_retries = 3
    retry_count = 0
    success = False
    
    while retry_count < max_retries and not success:
        try:
            print(f"🔄 尝试第 {retry_count + 1}/{max_retries} 次调用AI API...")
            run_log.append(f"🔄 尝试第 {retry_count + 1}/{max_retries} 次调用AI API...")
            
            # Create a new Word document
            doc = Document()
            chat = client.chat.create(model="grok-4")
            chat.append(system("You are Grok, a highly intelligent, helpful AI assistant."))
            chat.append(user(prompt))
            
            # 调用AI API
            print("⏳ 等待AI生成分析报告...")
            run_log.append("⏳ 等待AI生成分析报告...")
            response = chat.sample()
            
            # 验证响应内容
            if response and response.content and len(response.content.strip()) > 100:
                # 转换Markdown为Word格式
                print("📝 转换Markdown格式到Word...")
                run_log.append("📝 转换Markdown格式到Word...")
                
                doc = convert_markdown_to_word(response.content, doc)
                
                # Save file to the dated output directory
                output_filename = os.path.join(output_dir, f"{file_name}.docx")
                doc.save(output_filename)
                
                company_end_time = datetime.datetime.now()
                processing_time = (company_end_time - company_start_time).total_seconds()
                
                print(f"✅ 成功保存: {output_filename}")
                print(f"📄 内容长度: {len(response.content)} 字符")
                print(f"📝 已转换为格式化Word文档")
                print(f"⏱️  处理耗时: {processing_time:.1f} 秒")
                
                run_log.append(f"✅ 成功保存: {output_filename}")
                run_log.append(f"📄 内容长度: {len(response.content)} 字符")
                run_log.append(f"📝 已转换为格式化Word文档")
                run_log.append(f"⏱️  处理耗时: {processing_time:.1f} 秒")
                
                successful_reports.append({
                    'company': company,
                    'ticker': HKticker,
                    'filename': output_filename,
                    'content_length': len(response.content),
                    'processing_time': processing_time,
                    'completed_at': company_end_time.strftime('%H:%M:%S')
                })
                
                success = True
                
            else:
                print(f"⚠️  AI返回内容为空或过短，准备重试...")
                run_log.append(f"⚠️  AI返回内容为空或过短，准备重试...")
                retry_count += 1
                if retry_count < max_retries:
                    wait_time = 10 * retry_count  # 递增等待时间
                    print(f"⏸️  等待 {wait_time} 秒后重试...")
                    run_log.append(f"⏸️  等待 {wait_time} 秒后重试...")
                    time.sleep(wait_time)
                    
        except Exception as e:
            print(f"❌ 错误: {e}")
            run_log.append(f"❌ 错误: {e}")
            retry_count += 1
            if retry_count < max_retries:
                wait_time = 15 * retry_count  # 递增等待时间
                print(f"⏸️  等待 {wait_time} 秒后重试...")
                run_log.append(f"⏸️  等待 {wait_time} 秒后重试...")
                time.sleep(wait_time)
    
    if not success:
        print(f"🚫 {company} 处理失败，已达到最大重试次数")
        run_log.append(f"🚫 {company} 处理失败，已达到最大重试次数")
        
        # 创建一个错误文件记录
        error_filename = os.path.join(output_dir, f"ERROR - {file_name}.txt")
        with open(error_filename, 'w', encoding='utf-8') as f:
            f.write(f"处理失败\n公司: {company}\n股票代码: {HKticker}\n文件名: {file_name}\n时间: {datetime.datetime.now()}\n")
        print(f"📝 错误记录已保存: {error_filename}")
        run_log.append(f"📝 错误记录已保存: {error_filename}")
        
        failed_reports.append({
            'company': company,
            'ticker': HKticker,
            'error_file': error_filename,
            'failed_at': datetime.datetime.now().strftime('%H:%M:%S')
        })
    
    # 公司间等待时间
    if i < len(companies):  # 不是最后一家公司
        print(f"⏸️  等待 5 秒后处理下一家公司...")
        run_log.append(f"⏸️  等待 5 秒后处理下一家公司...")
        time.sleep(5)

# 计算总耗时
end_time = datetime.datetime.now()
total_time = (end_time - start_time).total_seconds()

print(f"\n🎉 所有公司处理完成!")
print(f"📁 文件保存位置: {output_dir}")
print(f"⏰ 完成时间: {end_time.strftime('%H:%M:%S')}")
print(f"⏱️  总耗时: {total_time:.1f} 秒 ({total_time/60:.1f} 分钟)")

run_log.append(f"\n🎉 所有公司处理完成!")
run_log.append(f"📁 文件保存位置: {output_dir}")
run_log.append(f"⏰ 完成时间: {end_time.strftime('%H:%M:%S')}")
run_log.append(f"⏱️  总耗时: {total_time:.1f} 秒 ({total_time/60:.1f} 分钟)")

# 生成运行报告Word文档
report_doc = Document()

# 添加标题
title = report_doc.add_heading('Investment Analysis Report Generation - Execution Report', 0)
title.alignment = 1  # 居中

# 添加总结段落
summary = report_doc.add_heading('📊 Execution Summary', level=1)

# 创建总结表格
summary_table = report_doc.add_table(rows=1, cols=2)
summary_table.style = 'Table Grid'

# 表格标题行
hdr_cells = summary_table.rows[0].cells
hdr_cells[0].text = 'Item'
hdr_cells[1].text = 'Result'

# 添加总结数据
summary_data = [
    ('Execution Date', today),
    ('Start Time', start_time.strftime('%H:%M:%S')),
    ('End Time', end_time.strftime('%H:%M:%S')),
    ('Total Duration', f"{total_time:.1f} seconds ({total_time/60:.1f} minutes)"),
    ('Total Companies to Process', str(len(companies))),
    ('Successful Reports Generated', str(len(successful_reports))),
    ('Failed Reports', str(len(failed_reports))),
    ('Success Rate', f"{len(successful_reports)/len(companies)*100:.1f}%")
]

for item, value in summary_data:
    row_cells = summary_table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = value

# 添加成功报告详情
if successful_reports:
    success_heading = report_doc.add_heading('✅ Successfully Generated Reports', level=1)
    
    success_table = report_doc.add_table(rows=1, cols=5)
    success_table.style = 'Table Grid'
    
    hdr_cells = success_table.rows[0].cells
    hdr_cells[0].text = 'Company Name'
    hdr_cells[1].text = 'Stock Code'
    hdr_cells[2].text = 'Completion Time'
    hdr_cells[3].text = 'Processing Time (sec)'
    hdr_cells[4].text = 'Content Length (chars)'
    
    for report in successful_reports:
        row_cells = success_table.add_row().cells
        row_cells[0].text = report['company']
        row_cells[1].text = report['ticker']
        row_cells[2].text = report['completed_at']
        row_cells[3].text = f"{report['processing_time']:.1f}"
        row_cells[4].text = str(report['content_length'])

# 添加失败报告详情
if failed_reports:
    failed_heading = report_doc.add_heading('❌ Failed Reports', level=1)
    
    failed_table = report_doc.add_table(rows=1, cols=3)
    failed_table.style = 'Table Grid'
    
    hdr_cells = failed_table.rows[0].cells
    hdr_cells[0].text = 'Company Name'
    hdr_cells[1].text = 'Stock Code'
    hdr_cells[2].text = 'Failed Time'
    
    for report in failed_reports:
        row_cells = failed_table.add_row().cells
        row_cells[0].text = report['company']
        row_cells[1].text = report['ticker']
        row_cells[2].text = report['failed_at']

# 添加详细运行日志
log_heading = report_doc.add_heading('📋 Detailed Execution Log', level=1)

for log_entry in run_log:
    report_doc.add_paragraph(log_entry)

# 保存运行报告
report_filename = os.path.join(output_dir, f"Execution Report - {today} - {end_time.strftime('%H%M%S')}.docx")
report_doc.save(report_filename)

# 断开LSEG连接
if lseg_connected:
    lseg_fetcher.disconnect()

print(f"📄 运行报告已保存: {report_filename}")
print(f"📈 成功率: {len(successful_reports)}/{len(companies)} ({len(successful_reports)/len(companies)*100:.1f}%)")