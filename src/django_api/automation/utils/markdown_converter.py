"""
Markdown to DOCX 转换模块

提供两种转换方式：
1. 内置转换器（默认）- 使用 python-docx，与原 Daily_Briefing_AI_Automation.py 一致
2. pypandoc（可选）- 使用 Pandoc 进行专业级转换，需要显式启用

使用方法：
    from automation.utils.markdown_converter import markdown_to_docx

    # 方式1: 直接保存文件（默认使用内置转换器）
    markdown_to_docx(markdown_text, output_path="/path/to/output.docx")

    # 方式2: 获取字节流
    docx_bytes = markdown_to_docx(markdown_text)

    # 方式3: 强制使用 Pandoc（如果需要更复杂的 Markdown 支持）
    markdown_to_docx(markdown_text, output_path="...", use_pandoc=True)

    # 方式4: 传入 Document 对象
    from docx import Document
    doc = Document()
    convert_markdown_to_word(markdown_text, doc)
    doc.save("/path/to/output.docx")
"""

from __future__ import annotations

import re
import tempfile
from io import BytesIO
from pathlib import Path
from typing import TYPE_CHECKING

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from observability import get_logger

if TYPE_CHECKING:
    pass

logger = get_logger(__name__)

# 导出的公共接口
__all__ = ["convert_markdown_to_word", "markdown_to_docx"]

# 检查 pypandoc 是否可用
_PYPANDOC_AVAILABLE = False
try:
    import pypandoc

    _PYPANDOC_AVAILABLE = True
except ImportError:
    pass  # pypandoc is optional, fallback to built-in converter

# 默认参考模板路径
_DEFAULT_REFERENCE_DOC = Path(__file__).resolve().parent.parent / "templates" / "reference.docx"


def markdown_to_docx(
    markdown_text: str,
    output_path: str | Path | None = None,
    title: str | None = None,
    reference_doc: str | Path | None = None,
    use_pandoc: bool = False,
) -> bytes | None:
    """
    将 Markdown 转换为 DOCX 格式。

    默认使用内置转换器（python-docx），与原 Daily_Briefing_AI_Automation.py 一致。
    可选使用 Pandoc 进行转换（需设置 use_pandoc=True）。

    Args:
        markdown_text: Markdown 格式的文本内容
        output_path: 输出文件路径，如果为 None 则返回字节流
        title: 文档标题（可选，仅 Pandoc 模式有效）
        reference_doc: 参考文档路径（可选，仅 Pandoc 模式有效）
        use_pandoc: 是否使用 Pandoc 转换（默认 False，使用内置转换器）

    Returns:
        如果 output_path 为 None，返回 DOCX 字节流；否则返回 None
    """
    # 如果明确不使用 Pandoc，或 Pandoc 不可用，使用内置转换器
    if not use_pandoc:
        return _builtin_markdown_to_docx(markdown_text, output_path)

    if not _PYPANDOC_AVAILABLE:
        logger.warning("pypandoc not available, using built-in converter")
        return _builtin_markdown_to_docx(markdown_text, output_path)

    try:
        # 构建 pandoc 额外参数
        extra_args = [
            "--standalone",  # 生成完整文档
            "--wrap=none",  # 不自动换行
        ]

        if title:
            extra_args.extend(["--metadata", f"title={title}"])

        # 使用参考模板（优先使用传入的，否则使用默认模板）
        ref_doc = reference_doc or _DEFAULT_REFERENCE_DOC
        if ref_doc and Path(ref_doc).exists():
            extra_args.extend(["--reference-doc", str(ref_doc)])

        if output_path:
            # 直接输出到文件
            pypandoc.convert_text(
                markdown_text,
                "docx",
                format="markdown",
                outputfile=str(output_path),
                extra_args=extra_args,
            )
            logger.info(
                "Markdown converted to DOCX using Pandoc",
                extra={"output_path": str(output_path), "content_length": len(markdown_text)},
            )
            return None
        # 输出到临时文件，然后读取字节
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name

        pypandoc.convert_text(
            markdown_text,
            "docx",
            format="markdown",
            outputfile=tmp_path,
            extra_args=extra_args,
        )

        with open(tmp_path, "rb") as f:
            docx_bytes = f.read()

        Path(tmp_path).unlink()  # 删除临时文件

        logger.info(
            "Markdown converted to DOCX bytes using Pandoc",
            extra={"bytes_size": len(docx_bytes), "content_length": len(markdown_text)},
        )
        return docx_bytes

    except Exception as e:
        logger.exception(
            "Pandoc conversion failed, using built-in converter", extra={"error": str(e)}
        )
        return _builtin_markdown_to_docx(markdown_text, output_path)


def _builtin_markdown_to_docx(
    markdown_text: str,
    output_path: str | Path | None = None,
) -> bytes | None:
    """
    内置转换器：使用 python-docx 进行转换。

    与原 Daily_Briefing_AI_Automation.py 中的转换逻辑一致。

    Args:
        markdown_text: Markdown 格式的文本内容
        output_path: 输出文件路径

    Returns:
        如果 output_path 为 None，返回 DOCX 字节流；否则返回 None
    """
    from docx import Document

    doc = Document()
    convert_markdown_to_word(markdown_text, doc)

    if output_path:
        doc.save(str(output_path))
        logger.info(
            "Markdown converted to DOCX",
            extra={"output_path": str(output_path), "content_length": len(markdown_text)},
        )
        return None
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def add_formatted_text(paragraph, text: str) -> None:
    """Add formatted text to paragraph with support for bold, italic, and inline code.

    Supports:
    - **bold text**
    - *italic text*
    - `inline code`
    - ***bold and italic***
    """
    # 先处理 inline code
    code_parts = re.split(r"(`[^`]+`)", text)

    for code_part in code_parts:
        if code_part.startswith("`") and code_part.endswith("`") and len(code_part) >= 2:
            run = paragraph.add_run(code_part[1:-1])
            run.font.name = "Courier New"
        else:
            # 使用非贪婪匹配处理加粗：**text**
            # 注意：使用 .+? 而不是 [^*]* 以支持更多情况
            bold_parts = re.split(r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*)", code_part)

            for bold_part in bold_parts:
                # 处理 ***bold and italic***
                if (
                    bold_part.startswith("***")
                    and bold_part.endswith("***")
                    and len(bold_part) >= 7
                ):
                    run = paragraph.add_run(bold_part[3:-3])
                    run.bold = True
                    run.italic = True
                # 处理 **bold**
                elif (
                    bold_part.startswith("**") and bold_part.endswith("**") and len(bold_part) >= 5
                ):
                    bold_content = bold_part[2:-2]
                    run = paragraph.add_run(bold_content)
                    run.bold = True
                else:
                    # 处理 *italic*
                    italic_parts = re.split(r"(\*[^*]+\*)", bold_part)

                    for italic_part in italic_parts:
                        if (
                            italic_part.startswith("*")
                            and italic_part.endswith("*")
                            and len(italic_part) >= 3
                        ):
                            run = paragraph.add_run(italic_part[1:-1])
                            run.italic = True
                        elif italic_part:
                            paragraph.add_run(italic_part)


def _strip_bold_markers(text: str) -> str:
    """Remove bold markers (**) from text for headings."""
    # Remove ***text*** -> text
    text = re.sub(r"\*\*\*([^*]+)\*\*\*", r"\1", text)
    # Remove **text** -> text
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    # Remove *text* -> text
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


def _is_bullet_list(line: str) -> bool:
    """Check if line is a bullet list item (not bold text starting with *)."""
    # Bullet list: starts with optional whitespace, then -, *, or +, then a space
    # But NOT ** (which is bold)
    stripped = line.lstrip()
    if stripped.startswith("**"):
        return False  # This is bold text, not a bullet
    return bool(re.match(r"^[ \t]*[-*+] ", line))


def convert_markdown_to_word(markdown_text, doc):
    """Convert Markdown text to Word document."""
    lines = markdown_text.split("\n")
    in_code_block = False
    in_table = False
    table_obj = None

    for line in lines:
        # Code block handling
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            continue

        # Table handling
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table = True
                cells = [cell.strip() for cell in line.strip().split("|")[1:-1]]
                table_obj = doc.add_table(rows=1, cols=len(cells))
                table_obj.style = "Table Grid"
                for i, cell_text in enumerate(cells):
                    # Add formatted text to table cell (supports bold/italic)
                    cell_paragraph = table_obj.rows[0].cells[i].paragraphs[0]
                    add_formatted_text(cell_paragraph, cell_text)
            elif (
                line.strip().replace("|", "").replace("-", "").replace(" ", "").replace(":", "")
                == ""
            ):  # Separator
                continue
            else:
                cells = [cell.strip() for cell in line.strip().split("|")[1:-1]]
                row_cells = table_obj.add_row().cells
                for i, cell_text in enumerate(cells):
                    if i < len(row_cells):
                        # Add formatted text to table cell (supports bold/italic)
                        cell_paragraph = row_cells[i].paragraphs[0]
                        add_formatted_text(cell_paragraph, cell_text)
            continue
        if in_table:
            in_table = False
            table_obj = None

        # Empty lines
        if not line.strip():
            doc.add_paragraph()
            continue

        # Headings - strip bold markers since headings are already bold
        if line.startswith("# ") and not line.startswith("## "):
            heading_text = _strip_bold_markers(line[2:].strip())
            doc.add_heading(heading_text, level=1)
        elif line.startswith("## ") and not line.startswith("### "):
            heading_text = _strip_bold_markers(line[3:].strip())
            doc.add_heading(heading_text, level=2)
        elif line.startswith("### ") and not line.startswith("#### "):
            heading_text = _strip_bold_markers(line[4:].strip())
            doc.add_heading(heading_text, level=3)
        elif line.startswith("#### "):
            heading_text = _strip_bold_markers(line[5:].strip())
            doc.add_heading(heading_text, level=4)

        # Bullet lists (but not bold text starting with *)
        elif _is_bullet_list(line):
            bullet_text = re.sub(r"^[ \t]*[-*+] ", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, bullet_text)

        # Numbered lists
        elif re.match(r"^[ \t]*\d+\. ", line):
            bullet_text = re.sub(r"^[ \t]*\d+\. ", "", line)
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, bullet_text)

        # Horizontal rule
        elif line.strip() in ["---", "***", "___"]:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run("─" * 50)
            run.font.color.rgb = RGBColor(128, 128, 128)

        # Blockquote
        elif line.strip().startswith("> "):
            quote_text = line.strip()[2:]
            p = doc.add_paragraph()
            p.left_indent = Inches(0.5)
            add_formatted_text(p, quote_text)

        # Normal text
        elif line.strip():
            p = doc.add_paragraph()
            add_formatted_text(p, line.strip())

    return doc
